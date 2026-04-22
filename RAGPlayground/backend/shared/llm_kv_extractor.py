"""
LLM-powered key-value extraction from IR Form-16 DOCX/PDF and SMAC PDF documents.

Hybrid approach for maximum accuracy + reasonable speed:
- DOCX: python-docx (reads XML directly — fast, accurate for structured tables)
- PDF:  Docling with OCR disabled (table structure detection without heavy OCR/image pipeline)
         Falls back to pypdf raw text if Docling fails or is unavailable.

Then sends batched table data to Ollama LLM for intelligent KV classification.
"""

import json
import re
import traceback
from typing import List, Dict, Tuple, Optional

from shared.ollama_client import ollama_chat


# ---------------------------------------------------------------------------
#  Nil value filter (same set as the legacy parser)
# ---------------------------------------------------------------------------
_NIL_VALUES = frozenset({
    "", "-", "\u2013", "\u2014", "nil", "-nil-", "n/a", "none",
    "not available", "not applicable", "na", ".", "..", "---",
    "nil.", "nil-", "-nil", "n.a.", "n.a",
})


# ---------------------------------------------------------------------------
#  LLM prompt templates
# ---------------------------------------------------------------------------

_IR_SYSTEM_PROMPT = """You are a precise document parser for Indian law enforcement IR Form-16 interrogation report tables.

Rules:
1. Each row may have 2, 3, or 4 columns. Column 1 is often a serial number or empty. Column 2 is usually the field name/description. Column 3 (and 4) are values.
2. When Column 1 is a number (e.g. "1", "2", "27"), it is the serial number for that field group.
3. When a row has empty first columns and text only in the last column, it is a CONTINUATION of the previous field's value. Combine the continuation text with the previous value.
4. When columns 3 and 4 have labels like "Present Address" and "Permanent Address", subsequent rows have two values (one for each). Create SEPARATE entries for each, e.g. "Village (Present Address)" and "Village (Permanent Address)".
5. Sub-labels like "(a)", "(b)", "(i)", "(ii)" are part of the field name. Keep them.
6. Skip rows that are pure headers like "Sl No", "Description", "Value" or section titles like "Personal Details", "Descriptive Roll", "Physical Description".
7. Skip fields where the value is only "-", "Nil", "N/A", "Not Available", or similar nil indicators.
8. For multi-line values that span multiple rows, combine them into one single value string.
9. If a cell contains both a field description and a value separated clearly, extract them as field_name and value.
10. Ignore duplicate/repeated columns (from merged cells) — use unique values only.

Return ONLY a JSON array. No explanation, no markdown code fences, no other text:
[{"serial_no": "1", "field_name": "(a) Full Name with alias", "value": "John Doe"}, ...]"""

_SMAC_SYSTEM_PROMPT = """You are a precise document parser for Indian law enforcement SMAC Log Report tables.

Rules:
1. The table typically has 3 columns: Serial Number | Field Name | Value.
2. Rows may use pipe (|) separators or multi-space separators.
3. Multi-line values belong to the most recent field name. Combine them into one value.
4. Common field names: TMS I.D., Originator, Date of Report, Theatre, Current Priority, Subject, Input, Grading, Has Attachment, Input Closed.
5. Skip header rows and empty rows.
6. Skip fields where the value is only "-", "Nil", "N/A", or similar nil indicators.

Return ONLY a JSON array. No explanation, no markdown code fences, no other text:
[{"serial_no": "1", "field_name": "TMS I.D.", "value": "TMS-2024-001"}, ...]"""


# ---------------------------------------------------------------------------
#  Docling singleton (OCR disabled, memory-optimized)
# ---------------------------------------------------------------------------
_converter = None


def _get_docling_converter():
    """Lazy-init Docling with OCR disabled and minimal memory footprint."""
    global _converter
    if _converter is not None:
        return _converter
    try:
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import (
            PdfPipelineOptions,
            TableFormerMode,
        )
        from docling.document_converter import DocumentConverter, PdfFormatOption
        from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend

        pipeline_options = PdfPipelineOptions()

        # Disable OCR — PDFs are text-based, not scanned images
        pipeline_options.do_ocr = False

        # Keep table structure detection (the main reason we use Docling)
        pipeline_options.do_table_structure = True
        pipeline_options.table_structure_options.do_cell_matching = True
        pipeline_options.table_structure_options.mode = TableFormerMode.ACCURATE

        # Disable all image/picture processing to save memory
        pipeline_options.generate_page_images = False
        pipeline_options.generate_picture_images = False
        pipeline_options.generate_table_images = False
        pipeline_options.do_picture_classification = False
        pipeline_options.do_picture_description = False
        pipeline_options.do_code_enrichment = False
        pipeline_options.do_formula_enrichment = False

        _converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(
                    pipeline_options=pipeline_options,
                    backend=PyPdfiumDocumentBackend,
                )
            }
        )
        print("[LLM-KV] Docling initialized (OCR disabled, table structure enabled)")
        return _converter
    except ImportError:
        print("[LLM-KV] Docling not installed, will use pypdf fallback for PDFs")
        return None
    except Exception as e:
        print(f"[LLM-KV] Docling init failed: {e}, will use pypdf fallback")
        return None


# ---------------------------------------------------------------------------
#  JSON response parsing (robust, handles LLM quirks)
# ---------------------------------------------------------------------------

def _validate_llm_kv_response(raw_text: str) -> List[Dict[str, str]]:
    """Parse LLM response into a list of KV dicts with robust error handling."""
    if not raw_text or not raw_text.strip():
        return []

    text = raw_text.strip()

    def _loads(s):
        return json.loads(s, strict=False)

    # Strategy 1: direct JSON parse
    try:
        parsed = _loads(text)
        if isinstance(parsed, list):
            return _filter_kv_list(parsed)
    except json.JSONDecodeError:
        pass

    # Strategy 2: strip markdown code fences first, then try JSON parse
    cleaned = re.sub(r'^```(?:json)?\s*[\r\n]*', '', text, flags=re.MULTILINE)
    cleaned = re.sub(r'[\r\n]*```\s*$', '', cleaned, flags=re.MULTILINE).strip()
    try:
        parsed = _loads(cleaned)
        if isinstance(parsed, list):
            return _filter_kv_list(parsed)
    except json.JSONDecodeError:
        pass

    # Strategy 3: extract [...] via regex
    match = re.search(r'\[.*\]', cleaned, re.DOTALL)
    if match:
        try:
            parsed = _loads(match.group())
            if isinstance(parsed, list):
                return _filter_kv_list(parsed)
        except json.JSONDecodeError:
            pass

    # Strategy 4: fix common JSON errors (trailing commas, single quotes)
    fixed = cleaned.replace("'", '"')
    fixed = re.sub(r',\s*([}\]])', r'\1', fixed)
    try:
        parsed = _loads(fixed)
        if isinstance(parsed, list):
            return _filter_kv_list(parsed)
    except json.JSONDecodeError:
        pass

    # Strategy 5: handle TRUNCATED JSON arrays (LLM ran out of tokens)
    truncated = cleaned if cleaned.lstrip().startswith('[') else fixed
    if '[' in truncated:
        arr_start = truncated.index('[')
        arr_text = truncated[arr_start:]
        last_brace = arr_text.rfind('}')
        if last_brace > 0:
            attempt = arr_text[:last_brace + 1].rstrip().rstrip(',') + '\n]'
            try:
                parsed = _loads(attempt)
                if isinstance(parsed, list):
                    print(f"[LLM-KV] Recovered {len(parsed)} items from truncated JSON response")
                    return _filter_kv_list(parsed)
            except json.JSONDecodeError:
                pass

    print(f"[LLM-KV] WARNING: Could not parse LLM response ({len(text)} chars)")
    return []


def _filter_kv_list(items: list) -> List[Dict[str, str]]:
    """Validate and filter KV entries."""
    result = []
    for item in items:
        if not isinstance(item, dict):
            continue
        fn = str(item.get("field_name", "")).strip()
        val = str(item.get("value", "")).strip()
        sno = str(item.get("serial_no", "")).strip()

        if not fn or not val:
            continue
        if val.lower().strip(".- ") in _NIL_VALUES:
            continue
        if re.match(r"^\d+\.?\s*$", fn):
            continue
        fn = fn.replace("**", "").strip()
        val = val.replace("**", "").strip()

        # Skip self-referencing entries (value == field_name)
        fn_norm = re.sub(r'[^a-z0-9]', '', fn.lower())
        val_norm = re.sub(r'[^a-z0-9]', '', val.lower())
        if fn_norm and val_norm and (fn_norm == val_norm or (fn_norm in val_norm and len(fn_norm) > 15)):
            continue

        result.append({"serial_no": sno, "field_name": fn, "value": val})
    return result


# ---------------------------------------------------------------------------
#  LLM call
# ---------------------------------------------------------------------------

def _llm_extract_kv_pairs(text: str, doc_type: str = "IR") -> List[Dict[str, str]]:
    """Send table/document text to the LLM and extract KV pairs."""
    system_prompt = _IR_SYSTEM_PROMPT if doc_type == "IR" else _SMAC_SYSTEM_PROMPT

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Extract all field_name/value pairs from this document data:\n\n{text}"},
    ]

    try:
        response = ollama_chat(messages, temperature=0.0)
        return _validate_llm_kv_response(response)
    except Exception as e:
        print(f"[LLM-KV] LLM call failed: {e}")
        return []


# ---------------------------------------------------------------------------
#  DOCX table extraction (python-docx — fast, reads XML directly)
# ---------------------------------------------------------------------------

def _extract_docx_table_rows(docx_path: str) -> Tuple[str, str]:
    """
    Extract all table rows from a DOCX using python-docx.
    Returns: (table_text, paragraph_text)
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(docx_path)
    lines = []
    row_idx = 0

    for table in doc.tables:
        for row in table.rows:
            raw_cells = [cell.text.strip().replace("\n", " ").replace("\r", " ") for cell in row.cells]
            # Deduplicate adjacent identical cells (merged cells produce duplicates)
            deduped = []
            for i, cell in enumerate(raw_cells):
                if i == 0 or cell != raw_cells[i - 1]:
                    deduped.append(cell)
            cell_text = " | ".join(deduped)
            lines.append(f"Row {row_idx}: {cell_text}")
            row_idx += 1

    table_text = f"Table data from the document ({row_idx} rows):\n" + "\n".join(lines) if lines else ""

    # Paragraph text (outside tables)
    para_parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            para_parts.append(t)
    paragraph_text = "\n".join(para_parts)

    return table_text, paragraph_text


# ---------------------------------------------------------------------------
#  PDF table extraction — Docling (preferred) or pypdf fallback
# ---------------------------------------------------------------------------

def _extract_pdf_tables_docling(pdf_path: str) -> Tuple[str, str]:
    """
    Extract tables from a PDF using Docling (OCR disabled).
    Returns: (table_text, paragraph_text)
    Falls back to pypdf if Docling fails.
    """
    converter = _get_docling_converter()
    if converter is None:
        return _extract_pdf_text_pypdf(pdf_path), ""

    try:
        print(f"[LLM-KV] Docling processing PDF: {pdf_path}")
        result = converter.convert(pdf_path)
        doc = result.document

        # Format tables for LLM
        lines = []
        row_idx = 0
        for table_idx, table in enumerate(doc.tables):
            try:
                df = table.export_to_dataframe(doc=doc)
                if df.empty or len(df) == 0:
                    continue
                # Emit column headers as a data row if they look like table data
                # (Docling sometimes treats the first data row as a header)
                col_cells = [str(c).strip().replace("\n", " ").replace("\r", " ") for c in df.columns]
                col_deduped = [col_cells[i] for i in range(len(col_cells)) if i == 0 or col_cells[i] != col_cells[i - 1]]
                col_text = " | ".join(col_deduped)
                # Only emit if the header looks like data (contains a serial number pattern)
                if col_deduped and any(c.rstrip(". ").isdigit() for c in col_deduped[:1]):
                    lines.append(f"Row {row_idx}: {col_text}")
                    row_idx += 1
                for r in range(len(df)):
                    raw_cells = [str(v).strip().replace("\n", " ").replace("\r", " ") for v in df.iloc[r]]
                    # Deduplicate adjacent identical cells (merged cells)
                    deduped = []
                    for i, cell in enumerate(raw_cells):
                        if i == 0 or cell != raw_cells[i - 1]:
                            deduped.append(cell)
                    cell_text = " | ".join(deduped)
                    lines.append(f"Row {row_idx}: {cell_text}")
                    row_idx += 1
            except Exception as e:
                print(f"[LLM-KV] Docling table {table_idx} failed: {e}")
                continue

        table_text = f"Table data from the document ({row_idx} rows):\n" + "\n".join(lines) if lines else ""

        # Paragraph text
        para_parts = []
        for text_item in doc.texts:
            t = str(text_item).strip()
            if t and not t.startswith("|") and not t.startswith("<!--"):
                para_parts.append(t)
        paragraph_text = "\n".join(para_parts)

        if not table_text:
            print("[LLM-KV] Docling found no tables, falling back to pypdf raw text")
            return _extract_pdf_text_pypdf(pdf_path), paragraph_text

        print(f"[LLM-KV] Docling extracted {row_idx} table rows")
        return table_text, paragraph_text

    except Exception as e:
        print(f"[LLM-KV] Docling failed: {e}, falling back to pypdf")
        return _extract_pdf_text_pypdf(pdf_path), ""


def _extract_pdf_text_pypdf(pdf_path: str) -> str:
    """Fallback: extract raw text from PDF using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(pdf_path)
    all_text = []
    for page_no, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
            if text:
                all_text.append(f"--- Page {page_no} ---\n{text}")
        except Exception as e:
            print(f"[LLM-KV] Page {page_no} extract error: {e}")
    return "\n\n".join(all_text)


# ---------------------------------------------------------------------------
#  Public API: DOCX extraction
# ---------------------------------------------------------------------------

def extract_kv_from_docx(docx_path: str) -> Tuple[List[str], List[Dict], str]:
    """
    Extract field_key/field_value pairs from an IR Form-16 DOCX.
    Uses python-docx for table extraction + LLM for KV classification.

    Returns: (units, metas, paragraph_text)
    """
    print(f"[LLM-KV] Extracting from DOCX: {docx_path}")
    table_text, paragraph_text = _extract_docx_table_rows(docx_path)

    if not table_text:
        print("[LLM-KV] No tables found in DOCX")
        return [], [], paragraph_text

    units: List[str] = []
    metas: List[Dict] = []
    total_llm_calls = 0

    # Split table text into chunks of ~6000 chars → typically 1-3 LLM calls
    chunks = _split_text_for_llm(table_text, max_chars=6000)
    print(f"[LLM-KV] Table text: {len(table_text)} chars → {len(chunks)} LLM call(s)")

    seen_fields = set()
    for chunk in chunks:
        kv_pairs = _llm_extract_kv_pairs(chunk, doc_type="IR")
        total_llm_calls += 1

        for kv in kv_pairs:
            fn = kv["field_name"]
            val = kv["value"]
            sno = kv.get("serial_no", "")

            dedup_key = f"{fn}||{val[:50]}"
            if dedup_key in seen_fields:
                continue
            seen_fields.add(dedup_key)

            units.append(f"{fn}: {val}")
            metas.append({"field_name": fn, "serial_no": sno, "section": ""})

    print(f"[LLM-KV] DOCX extraction complete: {len(units)} fields ({total_llm_calls} LLM calls)")
    return units, metas, paragraph_text


# ---------------------------------------------------------------------------
#  Public API: PDF extraction
# ---------------------------------------------------------------------------

def extract_kv_from_pdf(pdf_path: str, collection: str = "SMAC", total_pages: int = 1) -> Tuple[List[str], List[Dict]]:
    """
    Extract field_key/field_value pairs from a PDF.
    Uses Docling (OCR disabled) for table structure, falls back to pypdf.
    Then sends to LLM for KV classification.

    Args:
        pdf_path: path to the PDF file
        collection: "SMAC" or "IR" — determines which LLM prompt to use
        total_pages: total pages in the PDF (for page estimation in metadata)

    Returns: (units, metas)
    """
    try:
        from config import MAX_LLM_CALLS_PDF
    except ImportError:
        MAX_LLM_CALLS_PDF = 25

    doc_type = "IR" if collection == "IR" else "SMAC"
    print(f"[LLM-KV] Extracting from PDF ({doc_type} mode, {total_pages} pages): {pdf_path}")

    table_text, paragraph_text = _extract_pdf_tables_docling(pdf_path)

    if not table_text:
        print("[LLM-KV] No text extracted from PDF")
        return [], []

    units: List[str] = []
    metas: List[Dict] = []
    total_llm_calls = 0

    # Split into chunks and send to LLM
    chunks = _split_text_for_llm(table_text, max_chars=6000)
    chunks_to_process = min(len(chunks), MAX_LLM_CALLS_PDF)
    print(f"[LLM-KV] PDF text: {len(table_text)} chars → {len(chunks)} chunks, processing {chunks_to_process}")

    seen_fields = set()
    for i, chunk in enumerate(chunks[:MAX_LLM_CALLS_PDF]):
        print(f"[LLM-KV] Processing chunk {i+1}/{chunks_to_process}...")
        kv_pairs = _llm_extract_kv_pairs(chunk, doc_type=doc_type)
        total_llm_calls += 1

        # Estimate page from chunk position
        estimated_page = max(1, int((i / max(len(chunks), 1)) * total_pages) + 1)

        for kv in kv_pairs:
            fn = kv["field_name"]
            val = kv["value"]
            sno = kv.get("serial_no", "")

            dedup_key = f"{fn}||{val[:50]}"
            if dedup_key in seen_fields:
                continue
            seen_fields.add(dedup_key)

            units.append(f"{fn}: {val}")
            metas.append({"field_name": fn, "serial_no": sno, "page": estimated_page})

        print(f"[LLM-KV] Chunk {i+1}: extracted {len(kv_pairs)} fields (total so far: {len(units)})")

    if len(chunks) > MAX_LLM_CALLS_PDF:
        print(f"[LLM-KV] Processed {MAX_LLM_CALLS_PDF}/{len(chunks)} chunks (capped by MAX_LLM_CALLS_PDF)")

    print(f"[LLM-KV] PDF extraction complete: {len(units)} fields ({total_llm_calls} LLM calls)")
    return units, metas


# ---------------------------------------------------------------------------
#  Helper: split text into LLM-sized chunks
# ---------------------------------------------------------------------------

def _split_text_for_llm(text: str, max_chars: int = 6000) -> List[str]:
    """Split text into chunks at line boundaries."""
    if len(text) <= max_chars:
        return [text]

    chunks = []
    lines = text.split('\n')
    current_chunk = []
    current_len = 0

    for line in lines:
        if current_len + len(line) + 1 > max_chars and current_chunk:
            chunks.append('\n'.join(current_chunk))
            current_chunk = []
            current_len = 0
        current_chunk.append(line)
        current_len += len(line) + 1

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks
