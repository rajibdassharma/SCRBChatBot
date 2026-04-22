"""Document text and table extraction for PDF, DOCX, XLSX, CSV.
Includes OCR support for scanned PDFs via EasyOCR + PyMuPDF."""

import os
import re
from typing import List, Dict, Tuple, Optional


def extract_case_name(file_path: str) -> Optional[str]:
    """
    Extract case name from the immediate parent folder of the file.
    E.g., '/data/Bengaluru Blast Case/IR_John.docx' → 'Bengaluru Blast Case'
    """
    parent = os.path.basename(os.path.dirname(os.path.abspath(file_path)))
    if parent and parent not in (".", "..", "tmp", "temp", "uploads"):
        return parent
    return None


def extract_text(file_path: str, filename: str) -> Tuple[str, List[Dict]]:
    """
    Extract text from a document.
    Returns: (full_text, table_rows)
    - full_text: all text concatenated
    - table_rows: list of {row_idx, cells: [str]} for table rows
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif ext == ".xlsx":
        return _extract_xlsx(file_path)
    elif ext == ".csv":
        return _extract_csv(file_path)
    else:
        return "", []


def _extract_pdf(file_path: str) -> Tuple[str, List[Dict]]:
    """Extract text from PDF. Falls back to OCR if digital text is sparse."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    all_text = []
    for page in reader.pages:
        try:
            text = (page.extract_text() or "").strip()
            if text:
                all_text.append(text)
        except Exception:
            continue

    full_text = "\n\n".join(all_text)

    # If digital extraction yielded very little text, try OCR
    if len(full_text.strip()) < 100 and len(reader.pages) > 0:
        print(f"[DocumentLoader] Digital extraction sparse ({len(full_text)} chars), attempting OCR...")
        ocr_text = extract_pdf_ocr(file_path)
        if len(ocr_text) > len(full_text):
            return ocr_text, []

    return full_text, []


def extract_pdf_ocr(file_path: str) -> str:
    """OCR a scanned PDF using EasyOCR + PyMuPDF. Returns extracted text."""
    try:
        import fitz  # PyMuPDF
        import easyocr
        import numpy as np
        from PIL import Image
        import io
    except ImportError as e:
        print(f"[DocumentLoader] OCR dependencies not available: {e}")
        return ""

    print(f"[DocumentLoader] Running OCR on '{os.path.basename(file_path)}'...")

    # Initialize EasyOCR (English only, GPU if available)
    reader = easyocr.Reader(["en"], gpu=True, verbose=False)

    pdf_doc = fitz.open(file_path)
    all_text = []

    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        # Render page to image at 300 DPI for good OCR quality
        mat = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=mat)

        # Convert to numpy array for EasyOCR
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        img_np = np.array(img)

        # Run OCR
        results = reader.readtext(img_np, detail=1, paragraph=False)

        # Sort by vertical position then horizontal for reading order
        results.sort(key=lambda r: (r[0][0][1], r[0][0][0]))

        page_lines = []
        prev_y = -1
        current_line = []
        for bbox, text, conf in results:
            if conf < 0.2:  # Skip very low confidence
                continue
            top_y = bbox[0][1]
            # New line if Y position jumps significantly
            if prev_y >= 0 and abs(top_y - prev_y) > 15:
                if current_line:
                    page_lines.append(" ".join(current_line))
                current_line = []
            current_line.append(text)
            prev_y = top_y

        if current_line:
            page_lines.append(" ".join(current_line))

        page_text = "\n".join(page_lines)
        if page_text.strip():
            all_text.append(page_text)
            print(f"[DocumentLoader] OCR page {page_num + 1}/{len(pdf_doc)}: {len(page_text)} chars")

    pdf_doc.close()

    full_text = "\n\n".join(all_text)
    print(f"[DocumentLoader] OCR complete: {len(full_text)} chars from {len(pdf_doc)} pages")
    return full_text


def _extract_docx(file_path: str) -> Tuple[str, List[Dict]]:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)

    # Paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    para_text = "\n\n".join(paragraphs)

    # Tables
    table_rows = []
    row_idx = 0
    table_text_parts = []
    for table in doc.tables:
        for row in table.rows:
            raw_cells = [cell.text.strip() for cell in row.cells]
            # Deduplicate adjacent identical cells (merged cells)
            deduped = [raw_cells[i] for i in range(len(raw_cells))
                       if i == 0 or raw_cells[i] != raw_cells[i - 1]]
            table_rows.append({"row_idx": row_idx, "cells": deduped})
            table_text_parts.append(" | ".join(deduped))
            row_idx += 1

    full_text = para_text
    if table_text_parts:
        full_text += "\n\n" + "\n".join(table_text_parts)

    return full_text, table_rows


def _extract_xlsx(file_path: str) -> Tuple[str, List[Dict]]:
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    all_text = []
    table_rows = []
    row_idx = 0

    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            if any(cells):
                table_rows.append({"row_idx": row_idx, "cells": cells})
                all_text.append(" | ".join(cells))
                row_idx += 1

    wb.close()
    return "\n".join(all_text), table_rows


def _extract_csv(file_path: str) -> Tuple[str, List[Dict]]:
    import csv

    all_text = []
    table_rows = []
    row_idx = 0

    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            cells = [c.strip() for c in row]
            if any(cells):
                table_rows.append({"row_idx": row_idx, "cells": cells})
                all_text.append(" | ".join(cells))
                row_idx += 1

    return "\n".join(all_text), table_rows
