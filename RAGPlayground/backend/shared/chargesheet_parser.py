"""
Chargesheet Parser — LLM-based extraction from scanned PDF chargesheets.

Uses the local LLM (gemma3:12b) to extract structured data from OCR text:
  - Header fields (FIR No, Sections of Law, etc.)
  - Accused persons (Name, Age, Father's name, Address, etc.)
  - Pending/absconding persons
  - Brief Description of the case

The LLM handles OCR noise, broken text, and format variations that
regex parsers cannot reliably handle.
"""

import os
import re
import json
from typing import Dict, Any, List, Tuple

from shared.ollama_client import ollama_chat


# ── Encoding cleanup ────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = text.replace("\ufffd", "")
    text = re.sub(r"[ \t]{3,}", "  ", text)
    return text


# ── JSON parsing helpers ────────────────────────────────────────────────

def _extract_json(text: str) -> Any:
    """Extract JSON from LLM response — handles markdown fences, trailing commas, etc."""
    # Try direct parse
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try extracting from markdown code fence
    fence_match = re.search(r"```(?:json)?\s*\n?([\s\S]*?)```", text)
    if fence_match:
        try:
            return json.loads(fence_match.group(1).strip())
        except json.JSONDecodeError:
            pass

    # Try finding JSON array or object
    for pattern in [r"(\[[\s\S]*\])", r"(\{[\s\S]*\})"]:
        m = re.search(pattern, text)
        if m:
            raw = m.group(1)
            # Fix trailing commas
            raw = re.sub(r",\s*([}\]])", r"\1", raw)
            # Fix single quotes
            raw = raw.replace("'", '"')
            try:
                return json.loads(raw)
            except json.JSONDecodeError:
                pass

    return None


# ── LLM extraction prompts ──────────────────────────────────────────────

_HEADER_PROMPT = """You are extracting structured data from a chargesheet document (Indian police/NIA format).
This is OCR text from a scanned document. Read through any noise and garbled characters.

RULES:
1. ONLY extract information explicitly present in the text below.
2. Do NOT generate, invent, or hallucinate any information.
3. Respond ONLY in English. Do NOT use Hindi, Urdu, or any other language.
4. If a field is not found, skip it. Do NOT guess.

Extract the HEADER FIELDS (typically numbered 1-8 on the first page):
- Case title (STATE V/s who)
- Name of the Branch
- FIR Number
- Year
- Date
- Charge Sheet Number
- Sections of Law
- Type of Final Report
- Original or Supplementary
- Investigating Officer
- Complainant/Informant name

Return ONLY a JSON array like:
[{"field_name": "FIR Number", "value": "RC-01/2019/NIA/DLI"}, {"field_name": "Sections of Law", "value": "Under sections 120B, 121, 121A of IPC..."}, ...]

Do NOT include fields with empty or "Not Applicable" values.
Return ONLY the JSON array, nothing else. No explanation, no preamble.

TEXT:
"""

_ACCUSED_PROMPT = """You are extracting accused person details from a chargesheet document (Indian police/NIA format).
This is OCR text from a scanned document. Read through any garbled characters and OCR noise.

HOW TO IDENTIFY EACH ACCUSED:
- Each accused is marked with a number like A-1, A-2, A-3... (or (A-1), (A-2), etc.) before their details.
- After the A-number, there is a table/block with fields: Name, Age, Father's name, Sex, Nationality, Address, etc.
- Names often contain aliases separated by @ (e.g., "John Doe @ JD @ Johnny").
- The A-number may appear as "A-1", "(A-1)", "A1", "(A1)", or OCR variations like "A-l", "A -1".

RULES:
1. Look for EVERY A-number marker (A-1 through A-20 or more) and extract the accused person for each.
2. ONLY extract information that is explicitly present in the text below.
3. Do NOT generate, invent, or hallucinate any information not in the text.
4. Do NOT translate anything — keep all values in English exactly as they appear.
5. Respond ONLY in English. Do NOT use Hindi, Urdu, or any other language.
6. If a field is not found, set it to null. Do NOT guess.

Fields to extract per accused:
- accused_number (e.g., "A-1", "A-2"), name, age, fathers_name, sex, nationality, occupation, present_address, permanent_address, date_of_arrest, status, sections

Return ONLY a JSON array. Example:
[{"accused_number": "A-1", "name": "Tintin @ patla Tinu", "age": "25 years", "fathers_name": "Shri Lalu", "sex": "Male", "nationality": "Indian", "occupation": "Labourer", "present_address": "Murshidabad, West Bengal", "date_of_arrest": "26.08.2019", "status": "In judicial custody", "sections": "120B IPC, 18 UA(P) Act"}]

Return ONLY the JSON array, nothing else. No explanation, no preamble.

TEXT:
"""

_PENDING_PROMPT = """You are extracting details about accused/suspect persons whose investigation is pending.
This is from a chargesheet document (Indian police/NIA format), OCR text from scanned pages.

Look for a section titled "PARTICULARS OF THE ACCUSED AND SUSPECT PERSONS AGAINST WHOM THE INVESTIGATION IS PENDING" or similar.

Extract each person with their:
- Name
- Type: "accused", "suspect", or "absconder"
- Location/details (if mentioned)

Return ONLY a JSON array like:
[{"name": "Fareed Khan", "type": "accused", "details": "resident of Bangladesh"}, ...]

If this section is not found in the text, return an empty array: []
Return ONLY the JSON array, nothing else.

TEXT:
"""

_BRIEF_DESC_PROMPT = """You are extracting the "Brief Description of the Case" from a chargesheet document.
This is OCR text from a scanned document. Read through any garbled characters.

RULES:
1. ONLY extract text that is explicitly present in the document below.
2. Do NOT generate, invent, or hallucinate any content.
3. Respond ONLY in English. Do NOT use Hindi, Urdu, or any other language.
4. Clean up OCR noise but preserve the original content faithfully.

Look for a section titled "Brief Description" or "Brief Description about the Case" or similar.
Extract the ENTIRE brief description text.

If found, return the cleaned text as-is.
If not found, return exactly: NOT_FOUND

TEXT:
"""


# ── Main parser ─────────────────────────────────────────────────────────

def parse_chargesheet(file_path: str, filename: str, model: str = None) -> Dict[str, Any]:
    """
    Parse a chargesheet using LLM extraction.
    Chunks the text and sends each chunk to the LLM for structured extraction.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        full_text = _ocr_pdf(file_path)
    elif ext in (".docx", ".doc"):
        full_text = _extract_docx_text(file_path)
    else:
        return {"error": f"Unsupported file type: {ext}"}

    full_text = _clean_text(full_text)

    if len(full_text.strip()) < 50:
        return {"error": "No meaningful text extracted from chargesheet"}

    print(f"[ChargesheetParser] Extracted {len(full_text)} chars from '{filename}', starting LLM extraction...")

    # 1. Header fields — from first ~4000 chars
    header_fields = _llm_extract_header(full_text[:4000], model)

    # 2. Accused persons — chunk through the document
    accused_persons = _llm_extract_accused(full_text, model)

    # 3. Pending persons — from last ~6000 chars (usually near the end)
    pending_persons = _llm_extract_pending(full_text[-6000:], model)

    # 4. Brief description — from last ~8000 chars
    brief_desc = _llm_extract_brief_description(full_text[-8000:], model)

    # Build accused details text
    accused_text = ""
    if accused_persons:
        parts = []
        for p in accused_persons:
            line = f"Name: {p['person_name']}"
            if p.get("details"):
                line += f" | {p['details']}"
            parts.append(line)
        accused_text = "\n".join(parts)

    # Categorize pending persons
    absconder_text = ""
    suspect_text = ""
    if pending_persons:
        absconders = [p for p in pending_persons if p.get("person_type") in ("absconder", "accused")]
        suspects = [p for p in pending_persons if p.get("person_type") == "suspect"]
        absconder_text = "\n".join(f"{p['person_name']} ({p.get('details', '')})" for p in absconders)
        suspect_text = "\n".join(f"{p['person_name']} ({p.get('details', '')})" for p in suspects)

    result = {
        "header_fields": header_fields,
        "accused_persons": accused_persons,
        "accused_details_text": accused_text,
        "pending_persons": pending_persons,
        "absconder_details_text": absconder_text,
        "suspect_details_text": suspect_text,
        "brief_description": brief_desc,
        "full_text": full_text,
    }

    print(f"[ChargesheetParser] Parsed '{filename}': "
          f"{len(header_fields)} header fields, "
          f"{len(accused_persons)} accused, "
          f"{len(pending_persons)} pending persons, "
          f"brief_desc={len(brief_desc)} chars")

    return result


# ── LLM extraction functions ────────────────────────────────────────────

def _llm_extract_header(text: str, model: str = None) -> List[Dict[str, str]]:
    """Extract header fields using LLM."""
    print("[ChargesheetParser] Extracting header fields via LLM...")
    try:
        response = ollama_chat(
            [{"role": "user", "content": _HEADER_PROMPT + text}],
            temperature=0.0, model=model, max_tokens=2048,
        )
        data = _extract_json(response)
        if isinstance(data, list):
            fields = []
            for i, item in enumerate(data):
                fn = item.get("field_name", "").strip()
                val = item.get("value", "").strip()
                if fn and val:
                    fields.append({"serial_no": str(i + 1), "field_name": fn, "value": val})
            print(f"[ChargesheetParser] Header: extracted {len(fields)} fields")
            return fields
    except Exception as e:
        print(f"[ChargesheetParser] Header extraction failed: {e}")
    return []


_SINGLE_ACCUSED_PROMPT = """You are extracting ONE accused person's details from a chargesheet document.
This is OCR text from a scanned document. Read through any garbled characters.

RULES:
1. Extract the details of the accused person whose block starts in this text.
2. ONLY extract information explicitly present in the text. Do NOT invent anything.
3. Respond ONLY in English. Do NOT use Hindi, Urdu, or any other language.
4. If a field is not found, set it to null.

Fields: accused_number, name (including aliases with @), age, fathers_name, sex, nationality, occupation, present_address, permanent_address, date_of_arrest, status, sections

Return ONLY a JSON object (not array). Example:
{"accused_number": "A-1", "name": "Tintin @ patla Tinu", "age": "25 years", "fathers_name": "Shri Lalu", "sex": "Male", "nationality": "Indian", "occupation": "Labourer", "present_address": "Murshidabad, West Bengal", "date_of_arrest": "26.08.2019", "status": "In judicial custody", "sections": "120B IPC, 18 UA(P) Act"}

Return ONLY the JSON object, nothing else.

TEXT:
"""


_FIND_ACCUSED_COUNT_PROMPT = """You are analyzing a chargesheet document (Indian police/NIA format).
This is OCR text from a scanned document. Read through any garbled characters.

Your task: Find the TOTAL NUMBER of accused persons and their A-numbers.
Each accused is marked with A-1:-, A-2:-, A-3:- etc. (sometimes written as (A-1), A-1, A -1, or OCR variations).
The accused name usually appears right after the A-number marker.

RULES:
1. Count EVERY accused in the document. Do NOT miss any.
2. Return the accused number AND the name (including @ aliases) for each.
3. Respond ONLY in English.
4. ONLY list accused actually present in the text.

Return ONLY a JSON array like:
[{"a_number": "A-1", "name": "John Doe @ JD"}, {"a_number": "A-2", "name": "Tintin @ Tinu"}, ...]

Return ONLY the JSON array, nothing else.

TEXT:
"""


def _find_a_numbers_via_llm(full_text: str, model: str = None) -> List[Tuple[int, int]]:
    """Use LLM to find all A-number markers and their approximate positions in text."""
    print("[ChargesheetParser] Finding A-number markers via LLM...")

    # Send the full text in chunks to find all A-numbers
    # First try with a large chunk from the middle of the document (accused section)
    all_accused = []
    seen = set()

    # Process in overlapping large chunks to cover the whole document
    chunk_size = 8000
    step = 6000
    start = 0
    chunk_num = 0

    while start < len(full_text):
        chunk = full_text[start:start + chunk_size]
        if len(chunk.strip()) < 200:
            break
        chunk_num += 1

        try:
            response = ollama_chat(
                [{"role": "user", "content": _FIND_ACCUSED_COUNT_PROMPT + chunk}],
                temperature=0.0, model=model, max_tokens=2048,
            )
            data = _extract_json(response)
            if isinstance(data, list):
                for item in data:
                    a_num_str = item.get("a_number", "").strip()
                    # Normalize
                    m = re.search(r"(\d+)", a_num_str)
                    if m:
                        num = int(m.group(1))
                        if 1 <= num <= 30 and num not in seen:
                            seen.add(num)
                            all_accused.append({"num": num, "name": item.get("name", "")})
        except Exception as e:
            print(f"[ChargesheetParser] A-number scan chunk {chunk_num} failed: {e}")

        start += step

    if not all_accused:
        print("[ChargesheetParser] LLM found no A-number markers")
        return []

    all_accused.sort(key=lambda x: x["num"])
    labels = ["A-" + str(a["num"]) for a in all_accused]
    print(f"[ChargesheetParser] LLM found {len(all_accused)} accused: {labels}")

    # Find text positions — regex first, then name-based fallback for unlocated ones
    a_positions = []
    unlocated = []

    for acc in all_accused:
        num = acc["num"]

        patterns = [
            rf"A\s*-\s*{num}\s*[:\-]+",
            rf"\(\s*A\s*-\s*{num}\s*\)",
            rf"(?:^|\n)\s*A\s*-\s*{num}\s*[:\-)]",
        ]

        found = False
        for pat in patterns:
            m = re.search(pat, full_text)
            if m:
                a_positions.append((num, m.start()))
                found = True
                break

        if not found:
            unlocated.append(acc)

    # For unlocated ones, try finding by name in the text
    if unlocated:
        unloc_labels = ["A-" + str(a["num"]) for a in unlocated]
        print(f"[ChargesheetParser]   {len(unlocated)} not found by regex: {unloc_labels}, trying name search...")

        skip_words = {"the", "and", "shri", "smt", "mrs", "alias", "aged", "about", "son", "daughter", "wife"}
        for acc in unlocated:
            num = acc["num"]
            name = acc.get("name", "")
            if not name or len(name) < 4:
                print(f"[ChargesheetParser]   A-{num}: no name, skipping")
                continue

            name_words = [w for w in re.findall(r"[a-zA-Z]{3,}", name) if w.lower() not in skip_words]
            if len(name_words) < 2:
                print(f"[ChargesheetParser]   A-{num}: not enough name words, skipping")
                continue

            # Search for first two words appearing near each other
            w1, w2 = name_words[0], name_words[1]
            name_pat = re.compile(rf"{re.escape(w1)}.{{0,30}}{re.escape(w2)}", re.IGNORECASE)
            m = name_pat.search(full_text)
            if m:
                pos = max(0, m.start() - 100)
                a_positions.append((num, pos))
                print(f"[ChargesheetParser]   A-{num}: located by name '{w1}...{w2}' at pos {pos}")
            else:
                print(f"[ChargesheetParser]   A-{num}: could not locate in text")

    # Remove outliers: if LLM hallucinated A-13+ but doc only has A-1 to A-12,
    # detect gap and drop everything after it
    if a_positions:
        sorted_nums = sorted(set(n for n, _ in a_positions))
        cutoff = sorted_nums[-1]
        for i in range(1, len(sorted_nums)):
            if sorted_nums[i] - sorted_nums[i-1] > 2:
                cutoff = sorted_nums[i-1]
                break
        before = len(a_positions)
        a_positions = [(n, p) for n, p in a_positions if n <= cutoff]
        if len(a_positions) < before:
            print(f"[ChargesheetParser]   Dropped {before - len(a_positions)} outliers beyond A-{cutoff}")

    a_positions.sort(key=lambda x: x[1])
    verified = ["A-" + str(n) for n, _ in a_positions]
    print(f"[ChargesheetParser] Verified {len(a_positions)} accused positions: {verified}")
    return a_positions


_TARGETED_EXTRACT_PROMPT = """You are extracting details for ONE specific accused person from a chargesheet document.
This is OCR text from a scanned document. Read through any garbled characters.

I need the details for accused number {a_num} (marked as A-{a_num}:- in the document).

RULES:
1. Find ONLY the accused marked as A-{a_num} in this text.
2. ONLY extract information explicitly present in the text. Do NOT invent anything.
3. Respond ONLY in English. Do NOT use Hindi, Urdu, or any other language.
4. If you cannot find A-{a_num} in this text, return exactly: {{"not_found": true}}
5. If a field is not found, set it to null.

Fields: accused_number, name (including aliases with @), age, fathers_name, sex, nationality, occupation, present_address, permanent_address, date_of_arrest, status, sections

Return ONLY a JSON object. Example:
{{"accused_number": "A-{a_num}", "name": "Tintin @ patla Tinu", "age": "25 years", "fathers_name": "Shri Lalu", "sex": "Male", "nationality": "Indian", "occupation": "Labourer", "present_address": "Murshidabad, West Bengal", "date_of_arrest": "26.08.2019", "status": "In judicial custody", "sections": "120B IPC, 18 UA(P) Act"}}

Return ONLY the JSON object, nothing else.

TEXT:
"""


def _llm_extract_accused(full_text: str, model: str = None) -> List[Dict]:
    """Extract accused persons using a two-pass LLM approach:
    Pass 1: Discover all A-numbers and names
    Pass 2: For each accused, search through text chunks to extract full details
    """
    print("[ChargesheetParser] Extracting accused persons via LLM...")

    # Pass 1: Discover all accused
    all_accused = _discover_accused_list(full_text, model)

    if not all_accused:
        print("[ChargesheetParser] No accused found, falling back to chunk-based extraction")
        return _llm_extract_accused_chunked(full_text, model)

    # Remove outliers (gap detection)
    sorted_nums = sorted(a["num"] for a in all_accused)
    cutoff = sorted_nums[-1]
    for i in range(1, len(sorted_nums)):
        if sorted_nums[i] - sorted_nums[i - 1] > 2:
            cutoff = sorted_nums[i - 1]
            break
    all_accused = [a for a in all_accused if a["num"] <= cutoff]
    labels = ["A-" + str(a["num"]) for a in all_accused]
    print(f"[ChargesheetParser] Will extract {len(all_accused)} accused: {labels}")

    # Pass 2: For each accused, find the right chunk and extract details
    # Chunk the text with large overlap so every accused block is fully contained in at least one chunk
    chunk_size = 6000
    overlap = 2000
    chunks = []
    start = 0
    while start < len(full_text):
        end = start + chunk_size
        chunks.append(full_text[start:end])
        start = end - overlap

    all_persons = []

    for acc in all_accused:
        a_num = acc["num"]
        acc_name = acc.get("name", "")
        print(f"[ChargesheetParser]   Extracting A-{a_num} ({acc_name[:40]})...")

        # Try each chunk until we get a result for this accused
        extracted = False
        for chunk in chunks:
            prompt = _TARGETED_EXTRACT_PROMPT.format(a_num=a_num) + chunk

            try:
                response = ollama_chat(
                    [{"role": "user", "content": prompt}],
                    temperature=0.0, model=model, max_tokens=2048,
                )
                data = _extract_json(response)

                if isinstance(data, list) and data:
                    data = data[0]

                if not isinstance(data, dict):
                    continue

                if data.get("not_found"):
                    continue

                name = data.get("name", "").strip()
                if not name or len(name) < 3:
                    continue

                # Build details
                detail_parts = []
                for field in ["age", "fathers_name", "sex", "nationality", "occupation",
                              "present_address", "date_of_arrest", "status", "sections"]:
                    val = data.get(field)
                    if val and str(val).strip() and str(val).strip().lower() not in ("null", "none", "n/a", "not applicable", "not known"):
                        label = field.replace("_", " ").title()
                        detail_parts.append(f"{label}: {val}")

                person = {
                    "serial_no": f"A-{a_num}",
                    "person_name": name,
                    "person_type": "accused",
                    "details": "; ".join(detail_parts),
                }
                all_persons.append(person)
                print(f"[ChargesheetParser]   A-{a_num}: {name[:50]}")
                extracted = True
                break  # Got it, move to next accused

            except Exception as e:
                continue

        if not extracted:
            print(f"[ChargesheetParser]   A-{a_num}: could not extract from any chunk")

    print(f"[ChargesheetParser] Total accused extracted: {len(all_persons)}")
    return all_persons


def _discover_accused_list(full_text: str, model: str = None) -> List[Dict]:
    """Pass 1: Use LLM to discover all accused numbers and names."""
    all_accused = []
    seen = set()

    chunk_size = 8000
    step = 6000
    start = 0

    while start < len(full_text):
        chunk = full_text[start:start + chunk_size]
        if len(chunk.strip()) < 200:
            break

        try:
            response = ollama_chat(
                [{"role": "user", "content": _FIND_ACCUSED_COUNT_PROMPT + chunk}],
                temperature=0.0, model=model, max_tokens=2048,
            )
            data = _extract_json(response)
            if isinstance(data, list):
                for item in data:
                    a_num_str = item.get("a_number", "").strip()
                    m = re.search(r"(\d+)", a_num_str)
                    if m:
                        num = int(m.group(1))
                        if 1 <= num <= 30 and num not in seen:
                            seen.add(num)
                            all_accused.append({"num": num, "name": item.get("name", "")})
        except Exception as e:
            print(f"[ChargesheetParser] Discovery chunk failed: {e}")

        start += step

    all_accused.sort(key=lambda x: x["num"])
    labels = ["A-" + str(a["num"]) for a in all_accused]
    print(f"[ChargesheetParser] Discovery found {len(all_accused)} accused: {labels}")
    return all_accused


def _llm_extract_accused_chunked(full_text: str, model: str = None) -> List[Dict]:
    """Fallback: chunk-based extraction when no A-number markers found."""
    all_persons = []
    seen_names = set()
    chunk_size = 6000
    overlap = 1500
    text = full_text[800:]
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if len(chunk.strip()) >= 100:
            chunks.append(chunk)
        start = end - overlap

    for i, chunk in enumerate(chunks):
        try:
            response = ollama_chat(
                [{"role": "user", "content": _ACCUSED_PROMPT + chunk}],
                temperature=0.0, model=model, max_tokens=4096,
            )
            data = _extract_json(response)
            if isinstance(data, list):
                for item in data:
                    name = item.get("name", "").strip()
                    if not name or len(name) < 3:
                        continue
                    norm = re.sub(r"\s+", " ", name.lower()).split("@")[0].strip()
                    if norm in seen_names:
                        continue
                    seen_names.add(norm)
                    detail_parts = []
                    for field in ["age", "fathers_name", "sex", "nationality", "occupation",
                                  "present_address", "date_of_arrest", "status", "sections"]:
                        val = item.get(field)
                        if val and str(val).strip() and str(val).strip().lower() not in ("null", "none", "n/a", "not applicable", "not known"):
                            label = field.replace("_", " ").title()
                            detail_parts.append(f"{label}: {val}")
                    all_persons.append({
                        "serial_no": str(len(all_persons) + 1),
                        "person_name": name,
                        "person_type": "accused",
                        "details": "; ".join(detail_parts),
                    })
        except Exception as e:
            print(f"[ChargesheetParser] Chunk {i+1} failed: {e}")

    return all_persons


def _llm_extract_pending(text: str, model: str = None) -> List[Dict]:
    """Extract pending/absconding persons using LLM."""
    print("[ChargesheetParser] Extracting pending persons via LLM...")
    try:
        response = ollama_chat(
            [{"role": "user", "content": _PENDING_PROMPT + text}],
            temperature=0.0, model=model, max_tokens=2048,
        )
        data = _extract_json(response)
        if isinstance(data, list):
            persons = []
            for item in data:
                name = item.get("name", "").strip()
                ptype = item.get("type", "accused").strip().lower()
                if ptype in ("absconding", "wanted"):
                    ptype = "absconder"
                details = item.get("details", "").strip()
                if name and len(name) >= 2:
                    persons.append({
                        "person_name": name,
                        "person_type": ptype,
                        "details": details,
                    })
            print(f"[ChargesheetParser] Pending: extracted {len(persons)} persons")
            return persons
    except Exception as e:
        print(f"[ChargesheetParser] Pending extraction failed: {e}")
    return []


def _llm_extract_brief_description(text: str, model: str = None) -> str:
    """Extract brief description using LLM."""
    print("[ChargesheetParser] Extracting brief description via LLM...")
    try:
        response = ollama_chat(
            [{"role": "user", "content": _BRIEF_DESC_PROMPT + text}],
            temperature=0.0, model=model, max_tokens=4096,
        )
        response = response.strip()
        if response == "NOT_FOUND" or len(response) < 20:
            print("[ChargesheetParser] Brief description: not found")
            return ""
        # Remove any JSON wrapper if LLM wrapped it
        if response.startswith('"') and response.endswith('"'):
            response = response[1:-1]
        print(f"[ChargesheetParser] Brief description: {len(response)} chars")
        return response
    except Exception as e:
        print(f"[ChargesheetParser] Brief description extraction failed: {e}")
    return ""


# ── OCR / text extraction ───────────────────────────────────────────────

def _ocr_pdf(file_path: str) -> str:
    from shared.document_loader import extract_pdf_ocr, _extract_pdf

    digital_text, _ = _extract_pdf(file_path)
    if len(digital_text.strip()) > 200:
        print(f"[ChargesheetParser] Using digital text ({len(digital_text)} chars)")
        return digital_text

    ocr_text = extract_pdf_ocr(file_path)
    if ocr_text:
        return ocr_text
    return digital_text


def _extract_docx_text(file_path: str) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    parts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)
