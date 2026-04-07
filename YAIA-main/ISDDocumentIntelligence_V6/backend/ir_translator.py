"""
IR Translator — extracts Kannada narrative from IR DOCX and translates to English.

Separates tables (kept as-is) from free-form paragraphs (translated).
Generates output DOCX with original tables + English translation.
"""

import os
import re
import tempfile
from typing import Dict, List, Tuple, Any

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import TRANSLATE_MODEL
from ollama_client import ollama_chat


def _fix_encoding(text: str) -> str:
    if not text:
        return ""
    t = text.replace("\u2019", "'").replace("\u2018", "'")
    t = t.replace("\u201c", '"').replace("\u201d", '"')
    t = t.replace("\u2013", "-").replace("\u2014", "-")
    t = t.replace("\ufffd", "'")
    return t


def extract_tables_and_narrative(docx_path: str) -> Dict[str, Any]:
    """
    Parse an IR DOCX file and separate tables from narrative text.

    Returns:
        {
            "tables": list of table data (rows of cells),
            "narrative": str — free-form text after tables (Kannada),
            "filename": str,
        }
    """
    try:
        doc = DocxDocument(docx_path)
    except Exception as e:
        return {"error": f"Cannot open DOCX: {e}"}

    filename = os.path.basename(docx_path)

    # Extract tables
    tables = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [_fix_encoding(c.text).strip() for c in row.cells]
            # Deduplicate merged cells
            deduped = [cells[i] for i in range(len(cells))
                       if i == 0 or cells[i] != cells[i - 1]]
            table_rows.append(deduped)
        tables.append(table_rows)

    # Find which paragraphs belong to tables vs free text
    # python-docx exposes doc.element.body — tables and paragraphs are siblings
    # We collect all paragraphs that are NOT inside a table
    table_elements = set()
    for table in doc.tables:
        table_elements.add(table._element)

    narrative_parts = []
    past_last_table = False

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "tbl":
            past_last_table = True
            continue

        if tag == "p" and past_last_table:
            para_text = child.text or ""
            # Also get text from runs inside the paragraph
            runs_text = []
            for r in child.iter():
                if r.text:
                    runs_text.append(r.text)
            full_text = "".join(runs_text).strip()
            full_text = _fix_encoding(full_text)
            if full_text:
                narrative_parts.append(full_text)

    narrative = "\n\n".join(narrative_parts)

    print(f"[IR-Translator] {filename}: {len(tables)} tables, {len(narrative_parts)} narrative paragraphs, {len(narrative)} chars")

    return {
        "tables": tables,
        "narrative": narrative,
        "filename": filename,
    }


def translate_kannada_to_english(text: str) -> str:
    """Translate Kannada text to English using TranslateGemma."""
    if not text or not text.strip():
        return ""

    # Split into chunks if text is very long (TranslateGemma context limit)
    MAX_CHUNK = 3000
    if len(text) <= MAX_CHUNK:
        return _translate_chunk(text)

    # Split by paragraphs and translate in chunks
    paragraphs = text.split("\n\n")
    chunks = []
    current_chunk = ""

    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 > MAX_CHUNK:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = para
        else:
            current_chunk = current_chunk + "\n\n" + para if current_chunk else para

    if current_chunk:
        chunks.append(current_chunk)

    translated_chunks = []
    for i, chunk in enumerate(chunks):
        print(f"[IR-Translator] Translating chunk {i+1}/{len(chunks)} ({len(chunk)} chars)")
        translated = _translate_chunk(chunk)
        translated_chunks.append(translated)

    return "\n\n".join(translated_chunks)


def _translate_chunk(text: str) -> str:
    """Translate a single chunk of Kannada text."""
    prompt = (
        "Translate the following Kannada text to English. "
        "Return ONLY the English translation, nothing else. "
        "Do not include any preamble, explanation, or labels.\n\n"
        f"{text}"
    )

    try:
        result = ollama_chat(
            [{"role": "user", "content": prompt}],
            temperature=0.0,
            model=TRANSLATE_MODEL,
        )
        return result.strip()
    except Exception as e:
        print(f"[IR-Translator] Translation failed: {e}")
        return f"[Translation failed: {e}]"


def generate_translated_docx(
    tables: List[List[List[str]]],
    translated_text: str,
    original_filename: str,
) -> str:
    """
    Generate a new DOCX with original tables + translated English summary.

    Returns: path to the generated DOCX file.
    """
    doc = DocxDocument()

    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Translated IR — {original_filename}")
    run.bold = True
    run.font.size = Pt(14)

    # Add original tables
    for table_data in tables:
        if not table_data:
            continue

        # Determine max columns
        max_cols = max(len(row) for row in table_data) if table_data else 1

        doc_table = doc.add_table(rows=len(table_data), cols=max_cols)
        doc_table.style = "Table Grid"

        for r_idx, row in enumerate(table_data):
            for c_idx, cell_text in enumerate(row):
                if c_idx < max_cols:
                    doc_table.rows[r_idx].cells[c_idx].text = cell_text

        doc.add_paragraph()  # spacing

    # Add separator
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = separator.add_run("— Translated Summary —")
    run.bold = True
    run.font.size = Pt(12)

    # Add translated text
    for para in translated_text.split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)

    # Save to temp file
    tmp_dir = tempfile.mkdtemp()
    base = os.path.splitext(original_filename)[0]
    output_path = os.path.join(tmp_dir, f"{base}_translated.docx")
    doc.save(output_path)

    print(f"[IR-Translator] Generated: {output_path}")
    return output_path
