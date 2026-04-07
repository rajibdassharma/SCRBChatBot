"""
IR Document Parser — extracts structured fields from IR Form-16 DOCX/PDF files.

No LLM, no Docling. Pure Python table parsing.
Handles 3-column tables: [Serial No | Field Name | Value]
Skips free-form text sections after tables.

Returns a list of {serial_no, field_key, field_value} dicts.
"""

import re
import os
from typing import List, Dict, Tuple
from docx import Document as DocxDocument


# Values considered empty/nil
_NIL_VALUES = frozenset({
    "", "-", "–", "—", "nil", "-nil-", "n/a", "none",
    "not available", "not applicable", "na", ".", "..", "---",
})


def _fix_encoding(text: str) -> str:
    """Fix common encoding artifacts."""
    if not text:
        return ""
    t = text.replace("\u2019", "'").replace("\u2018", "'")
    t = t.replace("\u201c", '"').replace("\u201d", '"')
    t = t.replace("\u2013", "-").replace("\u2014", "-")
    t = t.replace("\ufffd", "'")
    return t


def _clean_label(text: str) -> str:
    """Clean field key/label: single line, collapsed whitespace."""
    t = _fix_encoding(text)
    t = t.replace("\r", " ").replace("\n", " ")
    t = re.sub(r"\s{2,}", " ", t)
    return t.strip()


def _clean_value(text: str) -> str:
    """Clean field value: preserve line breaks for multi-value fields."""
    t = _fix_encoding(text)
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse multiple blank lines but keep single line breaks
    t = re.sub(r"\n{3,}", "\n\n", t)
    # Clean up each line
    lines = [re.sub(r"\s{2,}", " ", ln).strip() for ln in t.split("\n")]
    # Remove empty lines at start/end
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    return "\n".join(lines)


def _get_row_cells(row) -> List[str]:
    """Extract and deduplicate adjacent identical cells (merged cells in DOCX).
    Returns raw text with encoding fixed but line breaks preserved."""
    raw = [_fix_encoding(c.text) for c in row.cells]
    return [raw[i].strip() for i in range(len(raw)) if i == 0 or raw[i].strip() != raw[i - 1].strip()]


def _is_serial(text: str) -> bool:
    """Check if text looks like a serial number.
    Matches: '1', '25', '130.', 'i', 'iv', 'xi', '(a)', '(p)', 'a)', 'a.', '1(a)'"""
    t = text.strip()
    # Numeric: 1, 25, 130.
    if re.match(r"^\d{1,3}\.?\s*$", t):
        return True
    # Roman numeral: i, ii, iii, iv, v, vi, vii, viii, ix, x, xi, xii
    if re.match(r"^[ivxlc]{1,6}\.?\s*$", t, re.IGNORECASE):
        return True
    # Letter with bracket/dot: (a), a), a., (p)
    if re.match(r"^\(?[a-z]\)?\.?\s*$", t, re.IGNORECASE):
        return True
    # Numeric + letter sub-item: 1(a), 58(p), 2a
    if re.match(r"^\d{1,3}\(?[a-z]\)?\s*$", t, re.IGNORECASE):
        return True
    return False


def _extract_serial(text: str) -> str:
    """Extract serial identifier from text. Returns the cleaned serial string."""
    t = text.strip().rstrip(". )")
    # Numeric: extract digits
    m = re.match(r"^(\d{1,3})", t)
    if m:
        # Check for sub-item: 58(p) → "58(p)"
        m2 = re.match(r"^(\d{1,3}\(?[a-z]\)?)", t, re.IGNORECASE)
        if m2:
            return m2.group(1)
        return m.group(1)
    # Roman numeral
    m = re.match(r"^([ivxlc]{1,6})", t, re.IGNORECASE)
    if m:
        return m.group(1).lower()
    # Letter: (a) → "a", a) → "a"
    m = re.match(r"^\(?([a-z])\)?", t, re.IGNORECASE)
    if m:
        return m.group(1).lower()
    return t


def _is_nil(value: str) -> bool:
    """Check if a value is considered nil/empty."""
    return value.lower().strip(".- ") in _NIL_VALUES


def parse_ir_docx(docx_path: str) -> Tuple[List[Dict[str, str]], str]:
    """
    Parse an IR Form-16 DOCX file.

    Returns:
        (fields, accused_name)
        - fields: List of {serial_no, field_key, field_value}
        - accused_name: Name extracted from the document
    """
    try:
        doc = DocxDocument(docx_path)
    except Exception as e:
        print(f"[IR-Parser] Cannot open DOCX '{docx_path}': {e}")
        return [], ""

    if not doc.tables:
        print(f"[IR-Parser] No tables found in {os.path.basename(docx_path)}")
        return [], ""

    fields: List[Dict[str, str]] = []
    accused_name = ""
    row_counter = 0
    _header_words = {"sl no", "sl no.", "description", "particulars", "serial no", "serial no."}

    for table in doc.tables:
        for row in table.rows:
            cells = _get_row_cells(row)

            if not cells or not any(c for c in cells):
                continue

            first = cells[0].strip()

            if _is_serial(first) and len(cells) >= 2:
                # Row with serial number: ['27.', 'Field Name', 'Value']
                serial = _extract_serial(first)
                field_key = _clean_label(cells[1])
                field_value = _clean_value(cells[2]) if len(cells) > 2 else ""

                if not field_key:
                    continue
                if _is_nil(field_value):
                    field_value = ""

                row_counter += 1
                fields.append({
                    "serial_no": serial,
                    "field_key": field_key,
                    "field_value": field_value,
                })

            elif len(cells) >= 2 and not first and cells[1].strip():
                # Empty first cell — check if second cell is a serial/sub-item marker
                second = cells[1].strip()
                if _is_serial(second) and len(cells) >= 3:
                    # Pattern: ['', '(v)', 'Field Name', 'Value']
                    serial = _extract_serial(second)
                    field_key = _clean_label(cells[2])
                    field_value = _clean_value(cells[3]) if len(cells) > 3 else ""

                    if not field_key:
                        continue
                    if _is_nil(field_value):
                        field_value = ""

                    row_counter += 1
                    fields.append({
                        "serial_no": serial,
                        "field_key": field_key,
                        "field_value": field_value,
                    })
                else:
                    # Pattern: ['', 'Field Name', 'Value']
                    field_key = _clean_label(second)

                    if not field_key:
                        continue
                    if field_key.lower().strip(".") in _header_words:
                        continue

                    field_value = _clean_value(cells[2]) if len(cells) > 2 else ""
                    if _is_nil(field_value):
                        field_value = ""

                    row_counter += 1
                    fields.append({
                        "serial_no": str(row_counter),
                        "field_key": field_key,
                        "field_value": field_value,
                    })

            elif len(cells) >= 2 and not _is_serial(first) and first:
                # Non-serial, non-empty first cell — could be:
                # 1. Header table with accused name
                # 2. Section header (e.g., "Financial Details") with sub-rows
                key = cells[0].lower()
                val = _clean_value(cells[1]) if len(cells) > 1 else ""

                if "name" in key and "accused" in key and val:
                    accused_name = val
                    fields.append({
                        "serial_no": "0",
                        "field_key": _clean_label(cells[0]),
                        "field_value": val,
                    })
                elif len(cells) >= 2:
                    # Treat as a field row — first cell is label, second is value
                    field_key = _clean_label(cells[0])
                    field_value = _clean_value(cells[1]) if len(cells) > 1 else ""
                    if len(cells) > 2 and cells[2].strip():
                        # 3-column row where first cell is a non-standard serial
                        field_key = _clean_label(cells[1])
                        field_value = _clean_value(cells[2])

                    if not field_key:
                        continue
                    if field_key.lower().strip(".") in _header_words:
                        continue
                    if _is_nil(field_value):
                        field_value = ""

                    row_counter += 1
                    fields.append({
                        "serial_no": str(row_counter),
                        "field_key": field_key,
                        "field_value": field_value,
                    })

    # Extract accused name from fields if not found in header
    if not accused_name:
        for f in fields:
            fk = f["field_key"].lower()
            if "name" in fk and "accused" in fk and f["field_value"]:
                accused_name = f["field_value"]
                break

    print(f"[IR-Parser] Extracted {len(fields)} fields from {len(doc.tables)} tables in {os.path.basename(docx_path)}")
    return fields, accused_name


def parse_ir_pdf(pdf_path: str) -> Tuple[List[Dict[str, str]], str]:
    """
    Parse an IR PDF file. Extracts text with pypdf, then parses
    for 3-column table patterns (serial | field | value).
    """
    from pypdf import PdfReader

    try:
        reader = PdfReader(pdf_path)
    except Exception as e:
        print(f"[IR-Parser] Cannot open PDF '{pdf_path}': {e}")
        return [], ""

    if not reader.pages:
        return [], ""

    fields: List[Dict[str, str]] = []
    accused_name = ""
    row_counter = 0

    # Patterns for 3-column rows
    # Pipe-separated: 1 | Field Name | Value
    pipe_re = re.compile(r"^(\d{1,3})\.?\s*\|\s*([^|]+?)\s*\|\s*(.*)$")
    # Space-separated: 1. Field Name    Value (2+ spaces between)
    space_re = re.compile(r"^(\d{1,3})[.)\s]\s+([A-Za-z][A-Za-z0-9 /(),\-\']{2,60}?)\s{2,}(.+)$")
    # Field without serial: Field Name    Value
    no_serial_re = re.compile(r"^([A-Za-z][A-Za-z0-9 /(),\-\']{2,60}?)\s{2,}(.+)$")

    for page in reader.pages:
        try:
            text = (page.extract_text() or "").strip()
        except Exception:
            continue

        if not text:
            continue

        lines = text.splitlines()

        for ln in lines:
            ln = ln.strip()
            if not ln:
                continue

            # Try pipe-separated
            m = pipe_re.match(ln)
            if m:
                serial = m.group(1)
                field_key = _clean_label(m.group(2))
                field_value = _clean_value(m.group(3))

                if not field_key:
                    continue
                if _is_nil(field_value):
                    field_value = ""

                row_counter += 1
                fields.append({
                    "serial_no": serial,
                    "field_key": field_key,
                    "field_value": field_value,
                })

                if "name" in field_key.lower() and "accused" in field_key.lower() and field_value:
                    accused_name = field_value
                continue

            # Try space-separated with serial
            m = space_re.match(ln)
            if m:
                serial = m.group(1)
                field_key = _clean_label(m.group(2))
                field_value = _clean_value(m.group(3))

                if not field_key:
                    continue
                if _is_nil(field_value):
                    field_value = ""

                row_counter += 1
                fields.append({
                    "serial_no": serial,
                    "field_key": field_key,
                    "field_value": field_value,
                })

                if "name" in field_key.lower() and "accused" in field_key.lower() and field_value:
                    accused_name = field_value
                continue

            # Try without serial number
            m = no_serial_re.match(ln)
            if m:
                field_key = _clean_label(m.group(1))
                field_value = _clean_value(m.group(2))

                if not field_key or len(field_key) < 3:
                    continue
                # Skip header-like lines
                if field_key.lower() in ("sl no", "description", "particulars", "serial no"):
                    continue
                if _is_nil(field_value):
                    field_value = ""

                row_counter += 1
                fields.append({
                    "serial_no": str(row_counter),
                    "field_key": field_key,
                    "field_value": field_value,
                })

                if "name" in field_key.lower() and "accused" in field_key.lower() and field_value:
                    accused_name = field_value

    # Extract accused name from fields if not found
    if not accused_name:
        for f in fields:
            fk = f["field_key"].lower()
            if "name" in fk and ("accused" in fk or fk == "name") and f["field_value"]:
                accused_name = f["field_value"]
                break

    print(f"[IR-Parser] PDF: Extracted {len(fields)} fields from {len(reader.pages)} pages in {os.path.basename(pdf_path)}")
    return fields, accused_name


def parse_ir_document(file_path: str, filename: str) -> Tuple[List[Dict[str, str]], str]:
    """
    Parse any IR document (DOCX or PDF).

    Returns:
        (fields, accused_name)
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".docx":
        return parse_ir_docx(file_path)
    elif ext == ".doc":
        # Try Word COM conversion on Windows
        try:
            docx_path = _convert_doc_to_docx(file_path)
            if docx_path:
                result = parse_ir_docx(docx_path)
                import shutil
                shutil.rmtree(os.path.dirname(docx_path), ignore_errors=True)
                return result
        except Exception as e:
            print(f"[IR-Parser] .doc conversion failed: {e}")
        return [], ""
    elif ext == ".pdf":
        return parse_ir_pdf(file_path)
    else:
        print(f"[IR-Parser] Unsupported format: {ext}")
        return [], ""


def _convert_doc_to_docx(doc_path: str):
    """Convert .doc to .docx using Word COM (Windows only)."""
    try:
        import win32com.client
        import pythoncom
        import tempfile

        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_path = os.path.abspath(doc_path)
        doc = word.Documents.Open(abs_path)
        tmp_dir = tempfile.mkdtemp()
        base = os.path.splitext(os.path.basename(doc_path))[0]
        docx_path = os.path.join(tmp_dir, base + ".docx")
        doc.SaveAs2(docx_path, FileFormat=12)
        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()
        if os.path.exists(docx_path):
            return docx_path
    except Exception as e:
        print(f"[IR-Parser] Word COM conversion failed: {e}")
    return None


# --- Test utility ---
if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python ir_parser.py <path_to_ir_file>")
        sys.exit(1)

    fpath = sys.argv[1]
    fname = os.path.basename(fpath)
    fields, name = parse_ir_document(fpath, fname)

    print(f"\nAccused: {name}")
    print(f"Fields: {len(fields)}\n")
    for f in fields:
        val_display = f["field_value"][:60] if f["field_value"] else "(empty)"
        print(f"  {f['serial_no']:>4} | {f['field_key'][:50]:50s} | {val_display}")
