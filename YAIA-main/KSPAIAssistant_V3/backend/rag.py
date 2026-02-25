"""
backend/rag.py

Robust RAG for PDF + Word (DOCX) + Excel (XLSX) + CSV documents, including table-heavy docs.

Key improvements vs your current version:
1) FIX: Do NOT discard short “field lines” (e.g., "Name Mohammed ...").
   - MIN_UNIT_LEN reduced (default 25).
2) PDF indexing stores per-page units + metadata (page, chunk_index).
3) XLSX indexing converts each row into key-value-like strings (header: value).
4) CSV indexing converts each row similarly.
5) ask_docs returns used_chunks metadata (doc_name, page/sheet, chunk_index)
6) clear_all_documents implemented safely by deleting and recreating collection.

Requirements:
    pip install pypdf chromadb python-docx openpyxl

Config expected (backend/config.py):
    CHROMA_PATH, EMBED_MODEL, PDF_MODEL
"""

import os
import re
import uuid
import csv
from typing import Dict, List, Any, Tuple, Optional

from pypdf import PdfReader
import chromadb

from docx import Document as DocxDocument
from openpyxl import load_workbook

from config import CHROMA_PATH, EMBED_MODEL, PDF_MODEL
from ollama_client import ollama_embed, ollama_chat


# -------------------------------------------------------------------
# Chroma persistent client + collection
# -------------------------------------------------------------------
_client = chromadb.PersistentClient(path=CHROMA_PATH)
_COLLECTION_NAME = "doc_chunks"
_collection = _client.get_or_create_collection(name=_COLLECTION_NAME)


# -------------------------------------------------------------------
# Utility: Safe chunking (guaranteed termination)
# -------------------------------------------------------------------
def chunk_text(text: str, chunk_size: int = 2000, overlap: int = 120) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []

    overlap = max(0, min(overlap, chunk_size - 1))
    chunks: List[str] = []
    n = len(text)
    start = 0

    while start < n:
        end = min(start + chunk_size, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= n:
            break

        next_start = end - overlap
        if next_start <= start:
            next_start = end
        start = next_start

    return chunks


# -------------------------------------------------------------------
# DOCX extraction: text + tables (structured rows)
# -------------------------------------------------------------------
def extract_text_and_tables_from_docx(docx_path: str) -> Tuple[str, List[str], Dict[str, Any]]:
    stats = {"type": "docx", "paragraphs": 0, "tables": 0, "table_rows": 0, "errors": []}
    doc = DocxDocument(docx_path)

    # Paragraph text
    paras = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            paras.append(t)
    stats["paragraphs"] = len(paras)
    full_text = "\n".join(paras)

    # Table rows as "Header: value | Header2: value2"
    table_rows: List[str] = []
    try:
        stats["tables"] = len(doc.tables)
        for ti, table in enumerate(doc.tables):
            rows = table.rows
            if not rows:
                continue

            header_cells = [c.text.strip() for c in rows[0].cells]
            has_headers = any(h for h in header_cells)

            for ri, row in enumerate(rows):
                cells = [c.text.strip().replace("\n", " ").strip() for c in row.cells]
                if not any(cells):
                    continue

                if has_headers and ri == 0:
                    continue

                if has_headers:
                    pairs = []
                    for h, v in zip(header_cells, cells):
                        h = h or f"Col{len(pairs)+1}"
                        v = v or ""
                        pairs.append(f"{h}: {v}")
                    row_text = " | ".join(pairs)
                else:
                    row_text = " | ".join([c for c in cells if c])

                row_text = f"[DOCX_TABLE {ti} ROW {ri}] {row_text}".strip()
                table_rows.append(row_text)

        stats["table_rows"] = len(table_rows)
    except Exception as e:
        stats["errors"].append(f"Table extract error: {e}")

    return full_text, table_rows, stats


# -------------------------------------------------------------------
# XLSX extraction: rows to key-value strings
# -------------------------------------------------------------------
def extract_rows_from_xlsx(xlsx_path: str, max_rows_per_sheet: int = 300) -> Tuple[List[str], Dict[str, Any]]:
    """
    Converts each row to:
      [SHEET Sheet1 ROW 12] ColA: val | ColB: val | ...
    Uses header row if present.
    """
    stats = {"type": "xlsx", "sheets": 0, "rows_emitted": 0, "errors": []}
    units: List[str] = []

    try:
        wb = load_workbook(xlsx_path, data_only=True)
        sheetnames = wb.sheetnames
        stats["sheets"] = len(sheetnames)

        for sname in sheetnames:
            ws = wb[sname]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            # Header row: first row with any text
            header = None
            for r in rows[:5]:
                if r and any(str(x).strip() for x in r if x is not None):
                    header = [str(x).strip() if x is not None else "" for x in r]
                    break

            start_idx = 1 if header else 0
            emitted = 0

            for i, r in enumerate(rows[start_idx:], start=start_idx + 1):
                if emitted >= max_rows_per_sheet:
                    break

                if not r or not any(x is not None and str(x).strip() for x in r):
                    continue

                pairs = []
                for ci, val in enumerate(r):
                    if val is None:
                        continue
                    v = str(val).strip()
                    if not v:
                        continue
                    key = header[ci] if header and ci < len(header) and header[ci] else f"Col{ci+1}"
                    pairs.append(f"{key}: {v}")

                if not pairs:
                    continue

                row_text = f"[SHEET {sname} ROW {i}] " + " | ".join(pairs)
                units.append(row_text)
                emitted += 1

            stats["rows_emitted"] += emitted

    except Exception as e:
        stats["errors"].append(str(e))

    return units, stats


# -------------------------------------------------------------------
# CSV extraction: rows to key-value strings
# -------------------------------------------------------------------
def extract_rows_from_csv(csv_path: str, max_rows: int = 500) -> Tuple[List[str], Dict[str, Any]]:
    stats = {"type": "csv", "rows_emitted": 0, "errors": []}
    units: List[str] = []

    try:
        with open(csv_path, "r", encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            all_rows = list(reader)

        if not all_rows:
            return [], stats

        header = [h.strip() for h in all_rows[0]]
        for i, r in enumerate(all_rows[1:], start=2):
            if stats["rows_emitted"] >= max_rows:
                break
            if not r or not any(str(x).strip() for x in r if x is not None):
                continue

            pairs = []
            for ci, val in enumerate(r):
                v = str(val).strip()
                if not v:
                    continue
                key = header[ci] if ci < len(header) and header[ci] else f"Col{ci+1}"
                pairs.append(f"{key}: {v}")

            if not pairs:
                continue

            units.append(f"[CSV ROW {i}] " + " | ".join(pairs))
            stats["rows_emitted"] += 1

    except Exception as e:
        stats["errors"].append(str(e))

    return units, stats


# -------------------------------------------------------------------
# Common indexing logic (caps + embedding + Chroma)
# -------------------------------------------------------------------
def _index_text_units(
    doc_id: str,
    filename: str,
    units: List[str],
    doc_type: str,
    extra_stats: Dict[str, Any],
    metas: Optional[List[dict]] = None,
    min_unit_len: int = 25,
) -> Dict[str, Any]:
    """
    units: list of strings to embed
    metas: optional list of metadata dicts aligned with units (same length)
    """

    # IMPORTANT FIX for your case:
    # Many form/table lines are SHORT (e.g., "Name Mohammed ..."), so 120 was too high.
    MIN_UNIT_LEN = max(0, min_unit_len)

    # Hard caps to avoid overload
    MAX_UNITS = 500

    cleaned_units: List[str] = []
    cleaned_metas: List[dict] = []

    for idx, u in enumerate(units):
        u2 = (u or "").strip()
        if len(u2) < MIN_UNIT_LEN:
            continue
        cleaned_units.append(u2)
        if metas and idx < len(metas):
            cleaned_metas.append(metas[idx])
        else:
            cleaned_metas.append({})

    total_before_cap = len(cleaned_units)
    cleaned_units = cleaned_units[:MAX_UNITS]
    cleaned_metas = cleaned_metas[:MAX_UNITS]

    print(f"[RAG] Indexing {doc_type.upper()}: {filename}")
    print(f"[RAG] Units before cap: {total_before_cap}, after cap: {len(cleaned_units)}")

    if not cleaned_units:
        return {"doc_id": doc_id, "doc_name": filename, "chunks": 0, "stats": extra_stats}

    ids: List[str] = []
    metadatas: List[dict] = []
    embeddings: List[list] = []
    documents: List[str] = []

    for i, (u, m) in enumerate(zip(cleaned_units, cleaned_metas)):
        if i % 25 == 0:
            print(f"[RAG] Embedding unit {i+1}/{len(cleaned_units)}")

        vec = ollama_embed(u, model=EMBED_MODEL)

        ids.append(f"{doc_id}_{i}")
        base_meta = {
            "doc_id": doc_id,
            "doc_name": filename,
            "doc_type": doc_type,
            "chunk_index": i,
        }
        base_meta.update(m or {})
        metadatas.append(base_meta)
        embeddings.append(vec)
        documents.append(u)

    _collection.add(ids=ids, embeddings=embeddings, documents=documents, metadatas=metadatas)

    return {"doc_id": doc_id, "doc_name": filename, "chunks": len(cleaned_units), "stats": extra_stats}


# -------------------------------------------------------------------
# Index PDF (page-aware + table-like lines)
# -------------------------------------------------------------------
def index_pdf(pdf_path: str, filename: str) -> Dict[str, Any]:
    doc_id = str(uuid.uuid4())

    reader = PdfReader(pdf_path)
    stats = {"type": "pdf", "pages": len(reader.pages), "extracted_chars": 0, "errors": []}

    units: List[str] = []
    metas: List[dict] = []

    MAX_PAGES = 120
    for page_no, page in enumerate(reader.pages[:MAX_PAGES], start=1):
        try:
            page_text = (page.extract_text() or "").strip()
        except Exception as e:
            stats["errors"].append(f"Page {page_no} extract error: {e}")
            page_text = ""

        if not page_text:
            continue

        stats["extracted_chars"] += len(page_text)

        # Chunk per-page text
        page_chunks = chunk_text(page_text, chunk_size=2000, overlap=140)

        # Keep table/form-like lines (THIS IS CRITICAL FOR FORMS)
        lines = [ln.strip() for ln in page_text.splitlines() if ln.strip()]

        # very permissive: ":" often appears in key/value lines
        tab_lines = [ln for ln in lines if (":" in ln or "|" in ln or ln.count("  ") >= 2)]
        tab_lines = tab_lines[:120]

        # Structured rows: S.No | Field Name | Field Value
        structured_units: List[str] = []
        structured_metas: List[dict] = []
        for ln in lines:
            # pipe-separated rows
            if "|" in ln:
                parts = [p.strip() for p in ln.split("|") if p.strip()]
                if len(parts) >= 3 and parts[0].replace(".", "").isdigit():
                    field_name = parts[1]
                    field_value = " | ".join(parts[2:])
                    structured_units.append(f"{field_name}: {field_value}")
                    structured_metas.append({"page": page_no, "field_name": field_name})
                    continue

            # colon-separated rows
            m = re.match(r"^(\d{1,3})[.)-]?\s*([^:]{2,80})\s*:\s*(.+)$", ln)
            if m:
                field_name = m.group(2).strip()
                field_value = m.group(3).strip()
                structured_units.append(f"{field_name}: {field_value}")
                structured_metas.append({"page": page_no, "field_name": field_name})
                continue

            # multi-space separated rows
            m2 = re.match(r"^(\d{1,3})[.)-]?\s+(.{2,80}?)\s{2,}(.+)$", ln)
            if m2:
                field_name = m2.group(2).strip()
                field_value = m2.group(3).strip()
                structured_units.append(f"{field_name}: {field_value}")
                structured_metas.append({"page": page_no, "field_name": field_name})

        page_units = page_chunks + tab_lines + structured_units

        for u in page_units:
            units.append(u)
            metas.append({"page": page_no})

        # Attach structured metadata (page + field_name) for structured units
        if structured_units:
            start_idx = len(units) - len(structured_units)
            for i, m in enumerate(structured_metas):
                if 0 <= start_idx + i < len(metas):
                    metas[start_idx + i].update(m)

    stats["units_total"] = len(units)
    return _index_text_units(
        doc_id,
        filename,
        units,
        doc_type="pdf",
        extra_stats=stats,
        metas=metas,
        min_unit_len=10,
    )


# -------------------------------------------------------------------
# Index DOCX
# -------------------------------------------------------------------
def index_docx(docx_path: str, filename: str) -> Dict[str, Any]:
    doc_id = str(uuid.uuid4())
    full_text, table_rows, stats = extract_text_and_tables_from_docx(docx_path)

    MAX_CHARS = 400_000
    if len(full_text) > MAX_CHARS:
        full_text = full_text[:MAX_CHARS]
        stats["truncated_to_chars"] = MAX_CHARS

    chunks = chunk_text(full_text, chunk_size=2200, overlap=140)
    stats["chunks_generated"] = len(chunks)

    units = chunks + table_rows
    return _index_text_units(doc_id, filename, units, doc_type="docx", extra_stats=stats)


# -------------------------------------------------------------------
# Index XLSX
# -------------------------------------------------------------------
def index_xlsx(xlsx_path: str, filename: str) -> Dict[str, Any]:
    doc_id = str(uuid.uuid4())
    units, stats = extract_rows_from_xlsx(xlsx_path)
    return _index_text_units(doc_id, filename, units, doc_type="xlsx", extra_stats=stats)


# -------------------------------------------------------------------
# Index CSV
# -------------------------------------------------------------------
def index_csv(csv_path: str, filename: str) -> Dict[str, Any]:
    doc_id = str(uuid.uuid4())
    units, stats = extract_rows_from_csv(csv_path)
    return _index_text_units(doc_id, filename, units, doc_type="csv", extra_stats=stats)


# -------------------------------------------------------------------
# Index any document
# -------------------------------------------------------------------
def index_document(file_path: str, filename: str) -> Dict[str, Any]:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return index_pdf(file_path, filename)
    if ext == ".docx":
        return index_docx(file_path, filename)
    if ext == ".xlsx":
        return index_xlsx(file_path, filename)
    if ext == ".csv":
        return index_csv(file_path, filename)

    return {"ok": False, "error": f"Unsupported document type: {ext}. Only PDF, DOCX, XLSX, CSV are supported."}


# -------------------------------------------------------------------
# Ask across documents
# -------------------------------------------------------------------
#def ask_docs(question: str, doc_ids: Optional[List[str]] = None, top_k: int = 10) -> Dict[str, Any]:
def ask_docs(question: str, doc_ids: Optional[List[str]] = None, top_k: int = 12) -> Dict[str, Any]:

    """
    If doc_ids is None -> search ALL indexed docs
    If doc_ids is [id1,id2] -> restrict to those docs (query per doc and merge)
    """
    qvec = ollama_embed(question, model=EMBED_MODEL)

    retrieved_docs: List[str] = []
    retrieved_meta: List[dict] = []

    if not doc_ids:
        res = _collection.query(
            query_embeddings=[qvec],
            n_results=top_k,
            include=["documents", "metadatas", "distances"],
        )
        retrieved_docs = (res.get("documents") or [[]])[0]
        retrieved_meta = (res.get("metadatas") or [[]])[0]
    else:
        per_doc_results: List[Tuple[str, dict]] = []
        per_doc_k = max(3, top_k // max(1, len(doc_ids)))

        for did in doc_ids:
            r = _collection.query(
                query_embeddings=[qvec],
                n_results=per_doc_k,
                where={"doc_id": did},
                include=["documents", "metadatas", "distances"],
            )
            docs = (r.get("documents") or [[]])[0]
            metas = (r.get("metadatas") or [[]])[0]
            for d, m in zip(docs, metas):
                per_doc_results.append((d, m))

        retrieved_docs = [x[0] for x in per_doc_results][:top_k]
        retrieved_meta = [x[1] for x in per_doc_results][:top_k]

    # Build context blocks
    used: List[dict] = []
    context_blocks: List[str] = []

    for d, m in zip(retrieved_docs, retrieved_meta):
        used.append({
            "doc_id": m.get("doc_id"),
            "doc_name": m.get("doc_name"),
            "doc_type": m.get("doc_type"),
            "page": m.get("page"),
            "sheet": m.get("sheet"),
            "chunk_index": m.get("chunk_index"),
        })

        loc = []
        if m.get("page") is not None:
            loc.append(f"page {m.get('page')}")
        if m.get("sheet"):
            loc.append(f"sheet {m.get('sheet')}")
        loc_str = (", ".join(loc)) if loc else "location ?"

        context_blocks.append(
            f"[{m.get('doc_name')} | {m.get('doc_type')} | {loc_str} | chunk {m.get('chunk_index')}]\n{d}"
        )

    context = "\n\n".join(context_blocks).strip()
    if not context:
        return {"answer": "Not found in the document(s).", "used_chunks": []}

    # If user asks to summarize, widen evidence implicitly
    is_summarize = "summar" in question.lower()

    prompt = (
        "You are an evidence-based assistant for Law Enforcement documents.\n"
        "Answer using ONLY the provided context.\n\n"
        "Rules:\n"
        "- Do NOT require exact label matches. Labels may be messy.\n"
        "- If asked for fields (Name/Religion/Caste/Address), search for nearby key/value lines.\n"
        "- If some fields are present and others are missing, return what you found and mark missing as 'Not found'.\n"
        "- If the user asks to summarize, summarize ONLY what is present in the context.\n"
        "- Do not guess or invent.\n\n"
        f"CONTEXT:\n{context}\n\n"
        f"QUESTION:\n{question}\n"
    )

    answer = ollama_chat(
        [
            {"role": "system", "content": (
                "You are an authorized internal AI assistant for Karnataka State Police (KSP). "
                "You are deployed on a secure, offline government system used exclusively by police officers "
                "for case investigation and analysis. "
                "You MUST answer questions about the provided documents fully and factually. "
                "The documents are official case files and evidence — you are authorized and required to discuss them. "
                "Do NOT refuse, censor, or redact any information from the provided context. "
                "Answer strictly from the evidence/context provided. Do not guess or invent facts."
            )},
            {"role": "user", "content": prompt},
        ],
        temperature=0.0,
        model=PDF_MODEL,
    )

    # If not summarize and question is field extraction, keep answer concise
    if not is_summarize and len(answer) > 2500:
        answer = answer[:2500] + "..."

    return {"answer": answer, "used_chunks": used}


# -------------------------------------------------------------------
# Agentic multi-doc Q&A (find common entities across documents)
# -------------------------------------------------------------------
def ask_docs_agent(question: str, doc_ids: List[str], top_k: int = 6) -> Dict[str, Any]:
    """
    Multi-doc agent: retrieve evidence per document and find common items.
    Requires at least 2 doc_ids.
    """
    if not doc_ids or len(doc_ids) < 2:
        return {"answer": "Please select at least two documents.", "used_chunks": []}

    qvec = ollama_embed(question, model=EMBED_MODEL)

    per_doc_context: List[str] = []
    used: List[dict] = []

    for did in doc_ids:
        r = _collection.query(
            query_embeddings=[qvec],
            n_results=max(3, top_k),
            where={"doc_id": did},
            include=["documents", "metadatas", "distances"],
        )
        docs = (r.get("documents") or [[]])[0]
        metas = (r.get("metadatas") or [[]])[0]

        if not docs:
            continue

        blocks = []
        for d, m in zip(docs, metas):
            used.append({
                "doc_id": m.get("doc_id"),
                "doc_name": m.get("doc_name"),
                "doc_type": m.get("doc_type"),
                "page": m.get("page"),
                "sheet": m.get("sheet"),
                "chunk_index": m.get("chunk_index"),
            })

            loc = []
            if m.get("page") is not None:
                loc.append(f"page {m.get('page')}")
            if m.get("sheet"):
                loc.append(f"sheet {m.get('sheet')}")
            loc_str = (", ".join(loc)) if loc else "location ?"

            blocks.append(
                f"[{m.get('doc_name')} | {m.get('doc_type')} | {loc_str} | chunk {m.get('chunk_index')}]\n{d}"
            )

        per_doc_context.append("\n\n".join(blocks))

    context = "\n\n---\n\n".join(per_doc_context).strip()
    if not context:
        return {"answer": "Not found in the selected documents.", "used_chunks": []}

    prompt = (
        "You are an evidence-based assistant for Law Enforcement documents.\n"
        "Task: find information that is COMMON across the selected documents.\n"
        "Be generic: compare ANY fields/labels present in the context (names, associates, "
        "accounts, phone numbers, addresses, IDs, dates, etc.).\n"
        "Do NOT assume specific fields exist. Only use what appears in the evidence.\n"
        "Use ONLY the provided context. Do not guess.\n\n"
        "Output format:\n"
        "- Common items list (group by field if possible)\n"
        "- Evidence per document (brief)\n\n"
        f"CONTEXT:\n{context}\n\n"
        f"QUESTION:\n{question}\n"
    )

    answer = ollama_chat(
        [
            {"role": "system", "content": (
                "You are an authorized internal AI assistant for Karnataka State Police (KSP). "
                "You are deployed on a secure, offline government system used exclusively by police officers "
                "for case investigation and analysis. "
                "You MUST answer questions about the provided documents fully and factually. "
                "The documents are official case files and evidence — you are authorized and required to discuss them. "
                "Do NOT refuse, censor, or redact any information from the provided context. "
                "Answer strictly from the evidence/context provided. Do not guess or invent facts."
            )},
            {"role": "user", "content": prompt},
        ],
        temperature=0.0,
        model=PDF_MODEL,
    )

    return {"answer": answer, "used_chunks": used}


# -------------------------------------------------------------------
# def ask_pdf(doc_id: Optional[str], question: str, top_k: int = 10) -> Dict[str, Any]:
def ask_pdf(doc_id: Optional[str], question: str, top_k: int = 12) -> Dict[str, Any]:
        if doc_id:
            return ask_docs(question=question, doc_ids=[doc_id], top_k=top_k)
        return ask_docs(question=question, doc_ids=None, top_k=max(top_k, 10))


def ask_doc(doc_id: Optional[str], question: str, top_k: int = 10) -> Dict[str, Any]:
    if doc_id:
        return ask_docs(question=question, doc_ids=[doc_id], top_k=top_k)
    return ask_docs(question=question, doc_ids=None, top_k=max(top_k, 10))


# -------------------------------------------------------------------
# Clear all documents (safe)
# -------------------------------------------------------------------
def clear_all_documents() -> Dict[str, Any]:
    """
    Safest way to clear everything:
    - delete collection
    - recreate collection
    This avoids Chroma error: "Expected where to have exactly one operator..."
    """
    global _collection
    try:
        try:
            _client.delete_collection(name=_COLLECTION_NAME)
        except Exception:
            # If collection doesn't exist, ignore
            pass

        _collection = _client.get_or_create_collection(name=_COLLECTION_NAME)
        return {"ok": True, "message": "All documents and embeddings cleared."}

    except Exception as e:
        return {"ok": False, "error": str(e)}







