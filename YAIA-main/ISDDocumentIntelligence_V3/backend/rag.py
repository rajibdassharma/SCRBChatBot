"""
backend/rag.py

Robust RAG for PDF + Word (DOCX) + Excel (XLSX) + CSV documents, including table-heavy docs.

Search Accuracy Features:
1) Hybrid Search: Vector (ChromaDB) + Keyword (BM25) with Reciprocal Rank Fusion
2) Multi-Query RAG: LLM generates 3 query variations for broader recall
3) LLM Re-ranking: Top candidates re-ranked by relevance before answer generation
4) Smart extraction: PDF forms/tables, DOCX tables, Excel/CSV row-to-key-value

Config expected (backend/config.py):
    CHROMA_PATH, EMBED_MODEL, PDF_MODEL,
    ENABLE_HYBRID_SEARCH, ENABLE_MULTI_QUERY, ENABLE_RERANKING
"""

import os
import re
import uuid
import csv
import json
from typing import Dict, List, Any, Tuple, Optional

from pypdf import PdfReader
import chromadb
from rank_bm25 import BM25Okapi

from docx import Document as DocxDocument
from openpyxl import load_workbook

from config import (
    CHROMA_PATH, EMBED_MODEL, PDF_MODEL,
    ENABLE_HYBRID_SEARCH, ENABLE_MULTI_QUERY, ENABLE_RERANKING,
    USE_LLM_PARSER,
)
from ollama_client import ollama_embed, ollama_embed_batch, ollama_chat
from structured_tables import store_smac_report, store_ir_report, search_fields, execute_sql_query, get_table_schema_description
from llm_kv_extractor import extract_kv_from_docx as _llm_extract_kv_from_docx
from llm_kv_extractor import extract_kv_from_pdf as _llm_extract_kv_from_pdf


# -------------------------------------------------------------------
# Chroma persistent client — multiple named collections
# -------------------------------------------------------------------
_client = chromadb.PersistentClient(path=CHROMA_PATH)
_col_cache: Dict[str, chromadb.Collection] = {}


def _to_chroma_name(col: str) -> str:
    """
    Map a logical collection name to a ChromaDB-compatible name.
    ChromaDB rules: 3–512 chars, [a-zA-Z0-9._-], must start and end with [a-zA-Z0-9].
    Examples: "IR" → "IR_db",  "SMAC" → "SMAC" (unchanged).
    """
    # Replace any character not in the allowed set with '_'
    name = re.sub(r"[^a-zA-Z0-9._-]", "_", col)
    # Ensure minimum length of 3
    if len(name) < 3:
        name = name + "_db"
    # Ensure start/end are alphanumeric
    if name and not name[0].isalnum():
        name = "c" + name
    if name and not name[-1].isalnum():
        name = name + "0"
    return name


def _get_col(col: str) -> chromadb.Collection:
    """Return (or create) a named ChromaDB collection. Builds BM25 on first access."""
    if col not in _col_cache:
        _col_cache[col] = _client.get_or_create_collection(name=_to_chroma_name(col))
        _rebuild_bm25(col)
    return _col_cache[col]


# -------------------------------------------------------------------
# BM25 In-Memory Index — per collection
# -------------------------------------------------------------------
_bm25: Dict[str, Dict] = {}


def _bs(col: str) -> Dict:
    """Return BM25 state dict for a collection, creating it if needed."""
    if col not in _bm25:
        _bm25[col] = {"docs": [], "metas": [], "index": None}
    return _bm25[col]


def _tokenize(text: str) -> List[str]:
    """Simple tokenizer: lowercase, split on non-alphanumeric, keep numbers."""
    return re.findall(r"[a-z0-9]+", text.lower())


def _rebuild_bm25(col: str):
    """Rebuild BM25 index from all documents in a ChromaDB collection."""
    state = _bs(col)
    try:
        collection = _col_cache[col]
        count = collection.count()
        if count == 0:
            state["docs"], state["metas"], state["index"] = [], [], None
            print(f"[BM25] No documents in '{col}', index empty")
            return

        all_data = collection.get(include=["documents", "metadatas"])
        state["docs"] = all_data.get("documents") or []
        state["metas"] = all_data.get("metadatas") or []

        corpus = [_tokenize(d) for d in state["docs"]]
        state["index"] = BM25Okapi(corpus) if corpus else None
        print(f"[BM25] Rebuilt index for '{col}' with {len(state['docs'])} chunks")
    except Exception as e:
        print(f"[BM25] Rebuild failed for '{col}': {e}")
        state["docs"], state["metas"], state["index"] = [], [], None


def _add_to_bm25(col: str, documents: List[str], metadatas: List[dict]):
    """Add new documents to the BM25 index for a collection."""
    state = _bs(col)
    state["docs"].extend(documents)
    state["metas"].extend(metadatas)
    corpus = [_tokenize(d) for d in state["docs"]]
    state["index"] = BM25Okapi(corpus) if corpus else None


def _clear_bm25(col: str):
    """Clear the BM25 index for a collection."""
    _bm25[col] = {"docs": [], "metas": [], "index": None}


def _bm25_search(col: str, query: str, top_k: int = 20, doc_ids: Optional[List[str]] = None) -> List[Tuple[str, dict]]:
    """Search BM25 index for a collection, return list of (document_text, metadata)."""
    state = _bs(col)
    if not state["index"] or not state["docs"]:
        return []

    tokenized_query = _tokenize(query)
    if not tokenized_query:
        return []

    scores = state["index"].get_scores(tokenized_query)

    scored = []
    for i, score in enumerate(scores):
        if score <= 0:
            continue
        if doc_ids and state["metas"][i].get("doc_id") not in doc_ids:
            continue
        scored.append((i, score))

    scored.sort(key=lambda x: x[1], reverse=True)
    results = [((state["docs"][i], state["metas"][i])) for i, _ in scored[:top_k]]
    return results


# -------------------------------------------------------------------
# Reciprocal Rank Fusion (merges Vector + BM25 results)
# -------------------------------------------------------------------
def _reciprocal_rank_fusion(
    vector_results: List[Tuple[str, dict]],
    bm25_results: List[Tuple[str, dict]],
    k: int = 60,
    top_n: int = 20,
) -> List[Tuple[str, dict]]:
    scores: Dict[str, float] = {}
    doc_map: Dict[str, Tuple[str, dict]] = {}

    for rank, (doc, meta) in enumerate(vector_results):
        key = f"{meta.get('doc_id', '')}_{meta.get('chunk_index', '')}_{doc[:100]}"
        scores[key] = scores.get(key, 0) + 1.0 / (k + rank + 1)
        doc_map[key] = (doc, meta)

    for rank, (doc, meta) in enumerate(bm25_results):
        key = f"{meta.get('doc_id', '')}_{meta.get('chunk_index', '')}_{doc[:100]}"
        scores[key] = scores.get(key, 0) + 1.0 / (k + rank + 1)
        doc_map[key] = (doc, meta)

    sorted_keys = sorted(scores.keys(), key=lambda x: scores[x], reverse=True)
    return [doc_map[k] for k in sorted_keys[:top_n]]


# -------------------------------------------------------------------
# Multi-Query Generation (LLM generates query variations)
# -------------------------------------------------------------------
def _generate_multi_queries(question: str) -> List[str]:
    """Generate 3 alternative phrasings of the question for broader retrieval coverage."""
    try:
        prompt = (
            "Generate 3 alternative search queries for finding relevant information "
            "in law enforcement case documents. Each query should use different keywords "
            "or approach the topic from a different angle.\n\n"
            f"Original question: {question}\n\n"
            "Return ONLY a JSON array of 3 strings. No explanation, no markdown fences.\n"
            'Example: ["arms and weapons training details", "firearms equipment used", "weapons training records"]'
        )

        response = ollama_chat(
            [{"role": "user", "content": prompt}],
            temperature=0.3,
            model=PDF_MODEL,
        )

        # Robust JSON extraction — use greedy match to get the full array
        text = response.strip()
        # Strip markdown fences
        text = re.sub(r'^```(?:json)?\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'\s*```\s*$', '', text, flags=re.MULTILINE).strip()

        # Find the outermost [...] using stack-based matching (handles nested/quoted brackets)
        start = text.find('[')
        if start != -1:
            depth = 0
            end = -1
            in_string = False
            escape_next = False
            for i, ch in enumerate(text[start:], start):
                if escape_next:
                    escape_next = False
                    continue
                if ch == '\\' and in_string:
                    escape_next = True
                    continue
                if ch == '"':
                    in_string = not in_string
                elif not in_string:
                    if ch == '[':
                        depth += 1
                    elif ch == ']':
                        depth -= 1
                        if depth == 0:
                            end = i + 1
                            break
            if end != -1:
                try:
                    queries = json.loads(text[start:end], strict=False)
                    if isinstance(queries, list) and len(queries) >= 1:
                        # Filter: discard queries that share no words with original question
                        # (prevents off-topic LLM hallucinations from polluting retrieval)
                        orig_words = set(re.findall(r"[a-zA-Z]{3,}", question.lower()))
                        valid = []
                        for q in queries:
                            if not isinstance(q, str) or not q.strip():
                                continue
                            q_words = set(re.findall(r"[a-zA-Z]{3,}", q.lower()))
                            if q_words & orig_words:  # must share at least one word
                                valid.append(q)
                        valid = valid[:3]
                        if valid:
                            print(f"[Multi-Query] Generated {len(valid)} variations")
                            return valid
                except json.JSONDecodeError:
                    pass

    except Exception as e:
        print(f"[Multi-Query] Failed: {e}")

    return []


# -------------------------------------------------------------------
# LLM Re-ranking (re-score top candidates by relevance)
# -------------------------------------------------------------------
def _rerank_with_llm(
    question: str,
    candidates: List[Tuple[str, dict]],
    top_n: int = 10,
) -> List[Tuple[str, dict]]:
    if len(candidates) <= top_n:
        return candidates

    try:
        passages = []
        for i, (doc, _meta) in enumerate(candidates):
            truncated = doc[:500] if len(doc) > 500 else doc
            passages.append(f"[{i}] {truncated}")

        passages_text = "\n\n".join(passages)

        prompt = (
            f"Question: {question}\n\n"
            f"Below are {len(candidates)} text passages from case documents. "
            f"Rank the top {top_n} most relevant passages for answering the question.\n\n"
            f"{passages_text}\n\n"
            f"Return ONLY a JSON array of the {top_n} most relevant passage indices, "
            f"ordered from most to least relevant.\n"
            f'Example: [5, 2, 8, 0, 11, 3, 7, 1, 9, 4]'
        )

        response = ollama_chat(
            [{"role": "user", "content": prompt}],
            temperature=0.0,
            model=PDF_MODEL,
        )

        match = re.search(r'\[.*?\]', response, re.DOTALL)
        if match:
            indices = json.loads(match.group())
            if isinstance(indices, list):
                reranked = []
                seen = set()
                for idx in indices:
                    if isinstance(idx, int) and 0 <= idx < len(candidates) and idx not in seen:
                        reranked.append(candidates[idx])
                        seen.add(idx)
                    if len(reranked) >= top_n:
                        break

                if reranked:
                    print(f"[Re-rank] Reordered {len(reranked)} candidates")
                    return reranked

    except Exception as e:
        print(f"[Re-rank] Failed: {e}, using original order")

    return candidates[:top_n]


# -------------------------------------------------------------------
# Utility: Safe chunking (guaranteed termination)
# -------------------------------------------------------------------
def chunk_text(text: str, chunk_size: int = 2000, overlap: int = 120) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []

    overlap = max(0, min(overlap, chunk_size - 1))
    chunks: List[str] = []
    n = len(text)
    start = 0

    while start < n:
        end = min(start + chunk_size, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= n:
            break

        next_start = end - overlap
        if next_start <= start:
            next_start = end
        start = next_start

    return chunks


# -------------------------------------------------------------------
# DOCX extraction: text + tables (structured rows)
# -------------------------------------------------------------------
def extract_text_and_tables_from_docx(docx_path: str) -> Tuple[str, List[str], Dict[str, Any]]:
    stats = {"type": "docx", "paragraphs": 0, "tables": 0, "table_rows": 0, "errors": []}
    doc = DocxDocument(docx_path)

    paras = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            paras.append(t)
    stats["paragraphs"] = len(paras)
    full_text = "\n".join(paras)

    table_rows: List[str] = []
    try:
        stats["tables"] = len(doc.tables)
        for ti, table in enumerate(doc.tables):
            rows = table.rows
            if not rows:
                continue

            header_cells = [c.text.strip() for c in rows[0].cells]
            has_headers = any(h for h in header_cells)

            for ri, row in enumerate(rows):
                cells = [c.text.strip().replace("\n", " ").strip() for c in row.cells]
                if not any(cells):
                    continue

                if has_headers and ri == 0:
                    continue

                if has_headers:
                    pairs = []
                    for h, v in zip(header_cells, cells):
                        h = h or f"Col{len(pairs)+1}"
                        v = v or ""
                        pairs.append(f"{h}: {v}")
                    row_text = " | ".join(pairs)
                else:
                    row_text = " | ".join([c for c in cells if c])

                row_text = f"[DOCX_TABLE {ti} ROW {ri}] {row_text}".strip()
                table_rows.append(row_text)

        stats["table_rows"] = len(table_rows)
    except Exception as e:
        stats["errors"].append(f"Table extract error: {e}")

    return full_text, table_rows, stats


# -------------------------------------------------------------------
# XLSX extraction: rows to key-value strings
# -------------------------------------------------------------------
def extract_rows_from_xlsx(xlsx_path: str, max_rows_per_sheet: int = 300) -> Tuple[List[str], Dict[str, Any]]:
    stats = {"type": "xlsx", "sheets": 0, "rows_emitted": 0, "errors": []}
    units: List[str] = []

    try:
        wb = load_workbook(xlsx_path, data_only=True)
        sheetnames = wb.sheetnames
        stats["sheets"] = len(sheetnames)

        for sname in sheetnames:
            ws = wb[sname]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            header = None
            for r in rows[:5]:
                if r and any(str(x).strip() for x in r if x is not None):
                    header = [str(x).strip() if x is not None else "" for x in r]
                    break

            start_idx = 1 if header else 0
            emitted = 0

            for i, r in enumerate(rows[start_idx:], start=start_idx + 1):
                if emitted >= max_rows_per_sheet:
                    break

                if not r or not any(x is not None and str(x).strip() for x in r):
                    continue

                pairs = []
                for ci, val in enumerate(r):
                    if val is None:
                        continue
                    v = str(val).strip()
                    if not v:
                        continue
                    key = header[ci] if header and ci < len(header) and header[ci] else f"Col{ci+1}"
                    pairs.append(f"{key}: {v}")

                if not pairs:
                    continue

                row_text = f"[SHEET {sname} ROW {i}] " + " | ".join(pairs)
                units.append(row_text)
                emitted += 1

            stats["rows_emitted"] += emitted

    except Exception as e:
        stats["errors"].append(str(e))

    return units, stats


# -------------------------------------------------------------------
# CSV extraction: rows to key-value strings
# -------------------------------------------------------------------
def extract_rows_from_csv(csv_path: str, max_rows: int = 500) -> Tuple[List[str], Dict[str, Any]]:
    stats = {"type": "csv", "rows_emitted": 0, "errors": []}
    units: List[str] = []

    try:
        with open(csv_path, "r", encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            all_rows = list(reader)

        if not all_rows:
            return [], stats

        header = [h.strip() for h in all_rows[0]]
        for i, r in enumerate(all_rows[1:], start=2):
            if stats["rows_emitted"] >= max_rows:
                break
            if not r or not any(str(x).strip() for x in r if x is not None):
                continue

            pairs = []
            for ci, val in enumerate(r):
                v = str(val).strip()
                if not v:
                    continue
                key = header[ci] if ci < len(header) and header[ci] else f"Col{ci+1}"
                pairs.append(f"{key}: {v}")

            if not pairs:
                continue

            units.append(f"[CSV ROW {i}] " + " | ".join(pairs))
            stats["rows_emitted"] += 1

    except Exception as e:
        stats["errors"].append(str(e))

    return units, stats


# -------------------------------------------------------------------
# Common indexing logic (caps + embedding + Chroma)
# -------------------------------------------------------------------
def _index_text_units(
    col: str,
    doc_id: str,
    filename: str,
    units: List[str],
    doc_type: str,
    extra_stats: Dict[str, Any],
    metas: Optional[List[dict]] = None,
    min_unit_len: int = 25,
) -> Dict[str, Any]:
    MIN_UNIT_LEN = max(0, min_unit_len)
    MAX_UNITS = 800

    cleaned_units: List[str] = []
    cleaned_metas: List[dict] = []

    for idx, u in enumerate(units):
        u2 = (u or "").strip()
        if len(u2) < MIN_UNIT_LEN:
            continue
        cleaned_units.append(u2)
        if metas and idx < len(metas):
            cleaned_metas.append(metas[idx])
        else:
            cleaned_metas.append({})

    total_before_cap = len(cleaned_units)
    cleaned_units = cleaned_units[:MAX_UNITS]
    cleaned_metas = cleaned_metas[:MAX_UNITS]

    print(f"[RAG:{col}] Indexing {doc_type.upper()}: {filename}")
    print(f"[RAG:{col}] Units before cap: {total_before_cap}, after cap: {len(cleaned_units)}")

    if not cleaned_units:
        return {"doc_id": doc_id, "doc_name": filename, "chunks": 0, "stats": extra_stats, "_chunks": []}

    print(f"[RAG:{col}] Batch-embedding {len(cleaned_units)} units...")
    embeddings = ollama_embed_batch(cleaned_units, model=EMBED_MODEL)
    print(f"[RAG:{col}] Embedding done, storing in ChromaDB...")

    ids: List[str] = []
    metadatas: List[dict] = []
    documents: List[str] = []

    for i, (u, m) in enumerate(zip(cleaned_units, cleaned_metas)):
        ids.append(f"{doc_id}_{i}")
        base_meta = {
            "doc_id": doc_id,
            "doc_name": filename,
            "doc_type": doc_type,
            "chunk_index": i,
        }
        base_meta.update(m or {})
        metadatas.append(base_meta)
        documents.append(u)

    _get_col(col).add(ids=ids, embeddings=embeddings, documents=documents, metadatas=metadatas)
    _add_to_bm25(col, documents, metadatas)

    return {"doc_id": doc_id, "doc_name": filename, "chunks": len(cleaned_units), "stats": extra_stats, "_chunks": cleaned_units}


# -------------------------------------------------------------------
# SMAC 3-column table parser (S.No | Field Name | Value)
# Handles multi-line values by accumulating continuation lines.
# -------------------------------------------------------------------
def _parse_smac_table_page_legacy(lines: List[str], page_no: int) -> Tuple[List[str], List[dict]]:
    """
    Legacy heuristic parser for SMAC/IR tabular PDF format.
    Supports two formats:
      3-column: Col1: Serial No | Col2: Field Name | Col3: Value
      2-column: Col1: Serial No | Col2: Full Text (all info in one column)
    Handles pipe-separated and multi-space-separated rows.
    Multi-line values are accumulated until the next numbered row.
    Returns: (units, metas) where each unit is "FieldName: Value".
    """
    units: List[str] = []
    metas: List[dict] = []
    current_field: Optional[str] = None
    current_serial: Optional[str] = None
    current_value_lines: List[str] = []

    # Pipe-separated:  1 | Gist | some value here
    pipe_re = re.compile(r"^(\d{1,3})\s*\|\s*([^|]+?)\s*\|\s*(.+)$")
    # Space-separated: 1  Gist  some value here  (2+ spaces as delimiter)
    space_re = re.compile(r"^(\d{1,3})[.)\s]\s+([A-Za-z][A-Za-z0-9 /()]{1,50}?)\s{2,}(.+)$")
    # 2-column pipe:   1 | full text here (only one pipe, no second column)
    pipe2_re = re.compile(r"^(\d{1,3})\s*\|\s*(.+)$")
    # Detects start of a new numbered row (to stop continuation accumulation)
    new_row_re = re.compile(r"^\d{1,3}[\s|.]")

    def flush_row():
        nonlocal current_field, current_serial, current_value_lines
        if current_field and current_value_lines:
            value = " ".join(current_value_lines).strip()
            if value:
                units.append(f"{current_field}: {value}")
                metas.append({"page": page_no, "field_name": current_field, "serial_no": current_serial or ""})
        current_field = None
        current_serial = None
        current_value_lines = []

    # ── Pass 1: try 3-column parsing ──────────────────────────────────
    for ln in lines:
        ln = ln.strip()
        if not ln:
            continue
        m = pipe_re.match(ln) or space_re.match(ln)
        if m:
            flush_row()
            current_serial = m.group(1).strip()
            current_field = m.group(2).strip()
            current_value_lines = [m.group(3).strip()]
        elif current_field:
            if not new_row_re.match(ln):
                current_value_lines.append(ln)
            else:
                flush_row()

    flush_row()

    if units:
        return units, metas

    # ── Pass 2: try 2-column parsing (Serial | Text) ─────────────────
    current_serial = None
    current_value_lines = []

    def flush_2col():
        nonlocal current_serial, current_value_lines
        if current_serial is not None and current_value_lines:
            text = " ".join(current_value_lines).strip()
            if text:
                # Try to extract "FieldName: Value" from the content
                colon_match = re.match(r"^([^:]{3,60}):\s+(.+)", text, re.DOTALL)
                if colon_match:
                    field = colon_match.group(1).strip()
                    val = colon_match.group(2).strip()
                    units.append(f"{field}: {val}")
                    metas.append({"page": page_no, "field_name": field, "serial_no": current_serial})
                else:
                    # No colon separator — use the text itself as content
                    # First few words become the field key for retrieval
                    words = text.split()
                    field = " ".join(words[:8]) if len(words) > 8 else text
                    units.append(f"{field}: {text}")
                    metas.append({"page": page_no, "field_name": field, "serial_no": current_serial})
        current_serial = None
        current_value_lines = []

    for ln in lines:
        ln = ln.strip()
        if not ln:
            continue
        m2 = pipe2_re.match(ln)
        if m2 and "|" not in m2.group(2):
            flush_2col()
            current_serial = m2.group(1).strip()
            current_value_lines = [m2.group(2).strip()]
        elif current_serial is not None:
            if not new_row_re.match(ln):
                current_value_lines.append(ln)
            else:
                flush_2col()

    flush_2col()
    return units, metas


# -------------------------------------------------------------------
# IR (Interrogation Report) DOCX table parser
# State-machine approach: iterates ALL tables in document order,
# classifies each row independently, and merges multi-row values.
# -------------------------------------------------------------------

# Row type constants
_IR_SKIP = "skip"
_IR_SECTION = "section"
_IR_FIELD_3 = "field_3"       # 3-col: [empty/sno, field_name, value]
_IR_FIELD_4 = "field_4"       # 4-col: address-type present/permanent split
_IR_FIELD_2 = "field_2"       # 2-col: [field_name, value]
_IR_CONTINUE = "continue"     # 2-col: ['', continuation_text]

_IR_HEADER_WORDS = frozenset({
    "sl no", "sl. no", "s.no", "s. no", "sno",
    "description", "particulars", "field", "serial no", "value",
})

_IR_SECTION_TITLE_RE = re.compile(
    r"(personal details|descriptive roll|physical description|"
    r"document details|financial details|communication details|"
    r"family members|criminal cases|apprehension|habits|"
    r"particulars of|role model|chronological|"
    r"part[\s\-]*[a-z]|section\s+\d)",
    re.IGNORECASE,
)

_IR_NIL_VALUES = frozenset({
    "-", "–", "—", "nil", "-nil-", "n/a", "none", "not available",
    "not applicable", "na", ".", "..", "---",
})


def _ir_get_row_cells(row) -> List[str]:
    """Extract and deduplicate cells from a DOCX table row."""
    raw = [
        c.text.strip().replace("\r", " ").replace("\n", " ").strip()
        for c in row.cells
    ]
    # Deduplicate adjacent identical cells (merged cells in DOCX)
    return [raw[i] for i in range(len(raw)) if i == 0 or raw[i] != raw[i - 1]]


def _ir_is_section_title(text: str) -> bool:
    """Heuristic: is this text a section header rather than a field value?"""
    t = text.strip()
    if not t:
        return False
    if _IR_SECTION_TITLE_RE.search(t):
        return True
    # Short text without sub-field markers, ending with period
    if len(t) < 60 and not re.search(r"\([a-z]\)|\([ivx]+\)", t):
        if t.endswith(".") and not any(ch.isdigit() for ch in t[:-1]):
            return True
    return False


def _ir_is_address_sub_header(col2: str, col3: str) -> bool:
    """Detect if a 4-col row is a sub-header defining column meanings."""
    markers = {"present", "permanent", "address", "temporary", "native"}
    c2 = col2.lower()
    c3 = col3.lower()
    return any(m in c2 for m in markers) and any(m in c3 for m in markers)


def _ir_classify_row(cells: List[str]) -> str:
    """Classify a DOCX table row into a row type for the IR parser."""
    non_empty = [c for c in cells if c.strip()]

    # Rule 1: nothing useful
    if not non_empty or len(cells) < 2:
        return _IR_SKIP

    # Rule 2: header row detection
    for c in cells:
        cleaned = c.strip().lower().strip(".:- ")
        if cleaned in _IR_HEADER_WORDS or (cleaned.startswith("sl") and len(cleaned) < 12):
            return _IR_SKIP

    first = cells[0].strip()
    is_serial = bool(re.match(r"^\d+\.?\s*$", first)) if first else False
    is_empty = first == ""

    # Rule 3: 4-column rows
    if len(cells) >= 4:
        return _IR_FIELD_4

    # Rule 4: 3-column rows
    if len(cells) == 3:
        second = cells[1].strip()
        third = cells[2].strip()

        if is_empty or is_serial:
            # Section header: serial + title, no value (or empty value)
            if is_serial and not third and _ir_is_section_title(second):
                return _IR_SECTION
            # Has both field name and value
            if second and third:
                return _IR_FIELD_3
            # Field name but empty value — could be section header or empty field
            if second and not third:
                return _IR_SECTION if _ir_is_section_title(second) else _IR_FIELD_3
        # First cell is text (not empty/serial) — treat as 3-col data
        if first and second:
            return _IR_FIELD_3

    # Rule 5: 2-column rows
    if len(cells) == 2:
        second = cells[1].strip()

        if is_empty:
            return _IR_CONTINUE if second else _IR_SKIP

        if is_serial:
            if second and _ir_is_section_title(second):
                return _IR_SECTION
            return _IR_CONTINUE if second else _IR_SKIP

        # Both cells are non-empty text
        if first and second:
            return _IR_FIELD_2

    return _IR_SKIP


def _ir_flush_field(
    field_name: str, value: str, serial_no: str, section: str,
    units: List[str], metas: List[dict],
):
    """Flush accumulated field state into the output lists."""
    fn = field_name.strip() if field_name else ""
    val = value.strip() if value else ""
    if not fn or not val:
        return
    # Skip nil/dash values
    if val.lower().strip(".- ") in _IR_NIL_VALUES:
        return
    # Skip pure-number field names
    if re.match(r"^\d+\.?\s*$", fn):
        return

    unit = f"{fn}: {val}"
    meta: Dict[str, str] = {"field_name": fn, "serial_no": serial_no}
    if section:
        meta["section"] = section
    units.append(unit)
    metas.append(meta)


def _parse_ir_table_from_docx_legacy(docx_path: str) -> Tuple[List[str], List[dict], str]:
    """
    Legacy heuristic parser for IR Form-16 DOCX. Row-by-row state machine
    that handles 2/3/4-column rows, multi-row value continuation,
    section headers, and address present/permanent splits.
    Returns: (units, metas, paragraph_text)
    """
    try:
        doc = DocxDocument(docx_path)
    except Exception as e:
        print(f"[RAG:IR] Cannot open as DOCX (file may be .doc renamed to .docx): {e}")
        return [], [], ""

    units: List[str] = []
    metas: List[dict] = []

    if not doc.tables:
        print("[RAG:IR] No tables found in DOCX")
        para_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        return units, metas, para_text

    # State machine variables
    current_section = ""
    current_field = None       # type: Optional[str]
    current_value = ""
    current_serial = ""
    four_col_headers = None  # Set when a 4-col sub-header is detected

    for table_idx, table in enumerate(doc.tables):
        for row in table.rows:
            cells = _ir_get_row_cells(row)
            row_type = _ir_classify_row(cells)

            if row_type == _IR_SKIP:
                continue

            elif row_type == _IR_SECTION:
                # Flush pending field
                _ir_flush_field(current_field, current_value, current_serial,
                                current_section, units, metas)
                current_field = None
                current_value = ""
                current_serial = ""
                four_col_headers = None
                # Extract section name
                section_text = cells[-1].strip() if cells[-1].strip() else (
                    cells[1].strip() if len(cells) > 1 else ""
                )
                section_text = re.sub(r"^\d+\.?\s*", "", section_text).strip().rstrip(".")
                if section_text:
                    current_section = section_text

            elif row_type == _IR_FIELD_3:
                # Flush pending field
                _ir_flush_field(current_field, current_value, current_serial,
                                current_section, units, metas)
                four_col_headers = None
                first = cells[0].strip()
                is_serial = bool(re.match(r"^\d+\.?\s*$", first)) if first else False
                if is_serial or first == "":
                    current_serial = first.rstrip(". ") if is_serial else ""
                    current_field = cells[1].strip()
                    current_value = cells[2].strip() if len(cells) > 2 else ""
                else:
                    # First cell is text — treat as field_name, join rest as value
                    current_serial = ""
                    current_field = first
                    current_value = " | ".join(
                        c.strip() for c in cells[1:] if c.strip()
                    )

            elif row_type == _IR_FIELD_4:
                # Flush pending field
                _ir_flush_field(current_field, current_value, current_serial,
                                current_section, units, metas)
                current_field = None
                current_value = ""

                first = cells[0].strip()
                is_serial = bool(re.match(r"^\d+\.?\s*$", first)) if first else False
                sno = first.rstrip(". ") if is_serial else ""
                label = cells[1].strip() if len(cells) > 1 else ""
                col2 = cells[2].strip() if len(cells) > 2 else ""
                col3 = cells[3].strip() if len(cells) > 3 else ""

                # Detect sub-header row (e.g. "Present Address | Permanent Address")
                if _ir_is_address_sub_header(col2, col3):
                    four_col_headers = (col2, col3)
                    continue

                if label:
                    if four_col_headers is not None:
                        # Dual-column mode (address section): emit two fields
                        if col2:
                            fn_left = f"{label} ({four_col_headers[0]})"
                            _ir_flush_field(fn_left, col2, sno, current_section, units, metas)
                        if col3:
                            fn_right = f"{label} ({four_col_headers[1]})"
                            _ir_flush_field(fn_right, col3, sno, current_section, units, metas)
                    else:
                        # No dual-column context: Col1=key, Col2=qualifier/value, Col3=value
                        if col2 and col3:
                            fn_combined = f"{label} - {col2}"
                            _ir_flush_field(fn_combined, col3, sno, current_section, units, metas)
                        elif col3:
                            _ir_flush_field(label, col3, sno, current_section, units, metas)
                        elif col2:
                            # col3 empty — col2 is the value
                            _ir_flush_field(label, col2, sno, current_section, units, metas)
                current_serial = ""

            elif row_type == _IR_FIELD_2:
                # Flush pending field
                _ir_flush_field(current_field, current_value, current_serial,
                                current_section, units, metas)
                four_col_headers = None
                current_serial = ""
                current_field = cells[0].strip()
                current_value = cells[1].strip()

            elif row_type == _IR_CONTINUE:
                continuation_text = cells[-1].strip()
                if continuation_text:
                    if current_field:
                        current_value += " " + continuation_text
                    elif current_section:
                        # Orphan continuation under a section — start section-level field
                        current_field = current_section
                        current_value = continuation_text
                        first = cells[0].strip()
                        current_serial = first.rstrip(". ") if (
                            first and re.match(r"^\d+\.?\s*$", first)
                        ) else ""
                    else:
                        # No active field or section yet — merged cell with
                        # section + field data. Store as introductory content.
                        current_field = "Status / Introduction"
                        current_value = continuation_text

    # Flush any remaining field after all tables
    _ir_flush_field(current_field, current_value, current_serial,
                    current_section, units, metas)

    print(f"[RAG:IR] Extracted {len(units)} field units from {len(doc.tables)} tables")

    # Extract paragraph text outside tables (narrative, officer notes, etc.)
    para_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

    return units, metas, para_text


# -------------------------------------------------------------------
# LLM-powered wrappers (with fallback to legacy heuristic parsers)
# -------------------------------------------------------------------

def _parse_ir_table_from_docx(docx_path: str) -> Tuple[List[str], List[dict], str]:
    """Parse IR Form-16 DOCX using Docling + LLM, falling back to legacy parser."""
    if USE_LLM_PARSER:
        try:
            units, metas, para = _llm_extract_kv_from_docx(docx_path)
            if units:
                print(f"[RAG:IR] Docling+LLM extraction: {len(units)} fields")
                return units, metas, para
            print("[RAG:IR] Docling+LLM returned 0 fields, falling back to heuristic parser")
        except Exception as e:
            print(f"[RAG:IR] Docling+LLM extraction failed: {e}")
            print("[RAG:IR] Falling back to heuristic parser")
    return _parse_ir_table_from_docx_legacy(docx_path)


def _parse_smac_table_page(lines: List[str], page_no: int) -> Tuple[List[str], List[dict]]:
    """Parse SMAC PDF page — delegates to legacy parser (PDF-level LLM handled in index_pdf)."""
    return _parse_smac_table_page_legacy(lines, page_no)


# -------------------------------------------------------------------
# Index PDF (page-aware + table-like lines)
# -------------------------------------------------------------------
def index_pdf(col: str, pdf_path: str, filename: str) -> Dict[str, Any]:
    doc_id = str(uuid.uuid4())

    reader = PdfReader(pdf_path)
    stats = {"type": "pdf", "pages": len(reader.pages), "extracted_chars": 0, "errors": []}

    units: List[str] = []
    metas: List[dict] = []

    # ── Try Docling+LLM for SMAC/IR PDFs (whole-file extraction) ──
    llm_pdf_units: List[str] = []
    llm_pdf_metas: List[dict] = []
    if col in ("SMAC", "IR") and USE_LLM_PARSER:
        try:
            llm_pdf_units, llm_pdf_metas = _llm_extract_kv_from_pdf(pdf_path, collection=col, total_pages=len(reader.pages))
            if llm_pdf_units:
                print(f"[RAG:PDF] Docling+LLM extraction: {len(llm_pdf_units)} fields from {filename}")
        except Exception as e:
            print(f"[RAG:PDF] Docling+LLM extraction failed: {e}, using per-page parser")

    MAX_PAGES = 120
    for page_no, page in enumerate(reader.pages[:MAX_PAGES], start=1):
        try:
            page_text = (page.extract_text() or "").strip()
        except Exception as e:
            stats["errors"].append(f"Page {page_no} extract error: {e}")
            page_text = ""

        if not page_text:
            continue

        stats["extracted_chars"] += len(page_text)

        page_chunks = chunk_text(page_text, chunk_size=2000, overlap=140)

        lines = [ln.strip() for ln in page_text.splitlines() if ln.strip()]

        if col in ("SMAC", "IR"):
            # Always add full-page text chunks for narrative content
            for u in page_chunks:
                units.append(u)
                metas.append({"page": page_no})

            # If Docling+LLM already extracted fields, skip per-page parsing
            if not llm_pdf_units:
                if col == "IR" and page_no == 1:
                    print(f"[RAG:IR] WARNING: LLM extraction unavailable, using basic per-page parser (may miss nested fields)")
                smac_units, smac_metas = _parse_smac_table_page(lines, page_no)
                if smac_units:
                    for u, sm in zip(smac_units, smac_metas):
                        units.append(u)
                        metas.append(sm)
        else:
            # Generic parsing for non-SMAC collections
            tab_lines = [ln for ln in lines if (":" in ln or "|" in ln or ln.count("  ") >= 2)]
            tab_lines = tab_lines[:120]

            structured_units: List[str] = []
            structured_metas: List[dict] = []
            for ln in lines:
                if "|" in ln:
                    parts = [p.strip() for p in ln.split("|") if p.strip()]
                    if len(parts) >= 3 and parts[0].replace(".", "").isdigit():
                        field_name = parts[1]
                        field_value = " | ".join(parts[2:])
                        structured_units.append(f"{field_name}: {field_value}")
                        structured_metas.append({"page": page_no, "field_name": field_name})
                        continue

                m = re.match(r"^(\d{1,3})[.)-]?\s*([^:]{2,80})\s*:\s*(.+)$", ln)
                if m:
                    field_name = m.group(2).strip()
                    field_value = m.group(3).strip()
                    structured_units.append(f"{field_name}: {field_value}")
                    structured_metas.append({"page": page_no, "field_name": field_name})
                    continue

                m2 = re.match(r"^(\d{1,3})[.)-]?\s+(.{2,80}?)\s{2,}(.+)$", ln)
                if m2:
                    field_name = m2.group(2).strip()
                    field_value = m2.group(3).strip()
                    structured_units.append(f"{field_name}: {field_value}")
                    structured_metas.append({"page": page_no, "field_name": field_name})

            page_units = page_chunks + tab_lines + structured_units
            for u in page_units:
                units.append(u)
                metas.append({"page": page_no})

            if structured_units:
                start_idx = len(units) - len(structured_units)
                for i, sm in enumerate(structured_metas):
                    if 0 <= start_idx + i < len(metas):
                        metas[start_idx + i].update(sm)

    # ── Append Docling+LLM extracted fields (if any) after page chunks ──
    if llm_pdf_units:
        for u, m in zip(llm_pdf_units, llm_pdf_metas):
            units.append(u)
            metas.append(m)

    stats["units_total"] = len(units)
    result = _index_text_units(col, doc_id, filename, units, doc_type="pdf", extra_stats=stats, metas=metas, min_unit_len=10)

    # ── Populate structured table (runs during indexing) ─────────────────
    if col in ("SMAC", "IR") and result.get("doc_id"):
        try:
            field_values = []
            narrative_parts = []
            for u, m in zip(units, metas):
                fn = m.get("field_name")
                if fn and ": " in u:
                    val = u.split(": ", 1)[1]
                    sno = m.get("serial_no", "")
                    field_values.append({"field_name": fn, "value": val, "serial_no": sno})
                elif not fn and len(u) > 50:
                    # Non-table text chunk (narrative/free-form) — collect it
                    narrative_parts.append(u)
            # Store narrative text as a special field so NL-to-SQL can find it
            if narrative_parts:
                narrative_text = " ".join(narrative_parts)[:4000]
                field_values.append({"field_name": "NARRATIVE_TEXT", "value": narrative_text, "serial_no": ""})
            if field_values:
                if col == "IR":
                    store_ir_report(result["doc_id"], filename, field_values)
                else:
                    store_smac_report(result["doc_id"], filename, field_values)
        except Exception as e:
            print(f"[StructuredTables] Store failed for {filename}: {e}")

    return result


# -------------------------------------------------------------------
# Index DOCX
# -------------------------------------------------------------------
def index_docx(col: str, docx_path: str, filename: str) -> Dict[str, Any]:
    doc_id = str(uuid.uuid4())

    if col == "IR":
        # IR Form-16: 3-column or 2-column table parser + paragraph text
        ir_units, ir_metas, ir_narrative = _parse_ir_table_from_docx(docx_path)
        if ir_units:
            # Combine table field units with narrative paragraph chunks
            all_units = list(ir_units)
            all_metas = list(ir_metas)
            narrative_chunks = []
            if ir_narrative.strip():
                narrative_chunks = chunk_text(ir_narrative, chunk_size=2000, overlap=140)
                for nc in narrative_chunks:
                    all_units.append(nc)
                    all_metas.append({"page": 1})

            print(f"[RAG:IR] Extracted {len(ir_units)} field units + {len(narrative_chunks)} narrative chunks from DOCX")
            stats = {"type": "docx", "mode": "ir_table", "field_units": len(ir_units), "narrative_chunks": len(narrative_chunks)}
            result = _index_text_units(
                col, doc_id, filename, all_units,
                doc_type="docx", extra_stats=stats, metas=all_metas, min_unit_len=5,
            )

            # ── Populate structured table (runs during indexing) ─────────
            if result.get("doc_id"):
                try:
                    field_values = []
                    for u, m in zip(ir_units, ir_metas):
                        fn = m.get("field_name")
                        if fn and ": " in u:
                            val = u.split(": ", 1)[1]
                            sno = m.get("serial_no", "")
                            field_values.append({"field_name": fn, "value": val, "serial_no": sno})
                    # Store narrative text as a special field for NL-to-SQL
                    if ir_narrative.strip():
                        field_values.append({"field_name": "NARRATIVE_TEXT", "value": ir_narrative.strip()[:4000], "serial_no": ""})
                    if field_values:
                        store_ir_report(result["doc_id"], filename, field_values)
                except Exception as e:
                    print(f"[StructuredTables] IR store failed for {filename}: {e}")

            return result
        print(f"[RAG:IR] No table rows found in DOCX, falling back to generic parsing")

    try:
        full_text, table_rows, stats = extract_text_and_tables_from_docx(docx_path)
    except Exception as e:
        # File may be a legacy .doc saved with .docx extension — try binary text extraction
        print(f"[RAG] DOCX open failed for '{filename}': {e} — attempting binary text fallback")
        raw_text = _extract_doc_text_fallback(docx_path)
        if not raw_text.strip():
            return {
                "ok": False,
                "error": (
                    f"'{filename}' could not be read as a DOCX file. "
                    "It may be a legacy .doc saved with a .docx extension. "
                    "Please re-save the file as proper .docx or convert to PDF."
                ),
            }
        chunks = chunk_text(raw_text, chunk_size=2000, overlap=140)
        fb_stats = {"type": "docx_binary_fallback", "chars": len(raw_text)}
        return _index_text_units(col, doc_id, filename, chunks, doc_type="docx", extra_stats=fb_stats)

    MAX_CHARS = 400_000
    if len(full_text) > MAX_CHARS:
        full_text = full_text[:MAX_CHARS]
        stats["truncated_to_chars"] = MAX_CHARS

    chunks = chunk_text(full_text, chunk_size=2200, overlap=140)
    stats["chunks_generated"] = len(chunks)

    units = chunks + table_rows
    return _index_text_units(col, doc_id, filename, units, doc_type="docx", extra_stats=stats)


# -------------------------------------------------------------------
# Index XLSX
# -------------------------------------------------------------------
def index_xlsx(col: str, xlsx_path: str, filename: str) -> Dict[str, Any]:
    doc_id = str(uuid.uuid4())
    units, stats = extract_rows_from_xlsx(xlsx_path)
    return _index_text_units(col, doc_id, filename, units, doc_type="xlsx", extra_stats=stats)


# -------------------------------------------------------------------
# Index CSV
# -------------------------------------------------------------------
def index_csv(col: str, csv_path: str, filename: str) -> Dict[str, Any]:
    doc_id = str(uuid.uuid4())
    units, stats = extract_rows_from_csv(csv_path)
    return _index_text_units(col, doc_id, filename, units, doc_type="csv", extra_stats=stats)


# -------------------------------------------------------------------
# .DOC (legacy binary Word) → convert to DOCX via Windows COM then index
# -------------------------------------------------------------------
def _convert_doc_to_docx_win32(doc_path: str) -> Optional[str]:
    """
    Convert a legacy .doc file to .docx using Word COM automation (Windows only).
    Returns the path to the temporary .docx file, or None on failure.
    """
    try:
        import win32com.client  # pywin32
        import pythoncom
        import tempfile

        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        abs_path = os.path.abspath(doc_path)
        doc = word.Documents.Open(abs_path)

        tmp_dir = tempfile.mkdtemp()
        base = os.path.splitext(os.path.basename(doc_path))[0]
        docx_path = os.path.join(tmp_dir, base + ".docx")

        # wdFormatXMLDocument = 12
        doc.SaveAs2(docx_path, FileFormat=12)
        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()

        if os.path.exists(docx_path):
            print(f"[DOC] Converted '{doc_path}' → '{docx_path}' via Word COM")
            return docx_path
    except Exception as e:
        print(f"[DOC] Word COM conversion failed: {e}")
    return None


def _extract_doc_text_fallback(doc_path: str) -> str:
    """
    Best-effort plain text extraction from .doc binary (no Word/LibreOffice).
    Extracts readable ASCII runs. Quality is limited.
    """
    try:
        with open(doc_path, "rb") as f:
            content = f.read()
        # Extract runs of printable ASCII (≥6 chars) to reduce noise
        runs = re.findall(rb"[\x20-\x7e]{6,}", content)
        lines = []
        for run in runs:
            text = run.decode("ascii", errors="ignore").strip()
            # Filter out lines that are mostly non-word characters (binary noise)
            word_chars = sum(1 for c in text if c.isalpha() or c.isdigit() or c == " ")
            if text and word_chars / max(len(text), 1) > 0.5:
                lines.append(text)
        return "\n".join(lines)
    except Exception as e:
        print(f"[DOC] Binary fallback extraction failed: {e}")
        return ""


# -------------------------------------------------------------------
# Index any document
# -------------------------------------------------------------------
def index_document(file_path: str, filename: str, collection_name: str = "SMAC") -> Dict[str, Any]:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return index_pdf(collection_name, file_path, filename)
    if ext == ".docx":
        return index_docx(collection_name, file_path, filename)
    if ext == ".doc":
        # Try Word COM conversion to DOCX first (Windows with MS Word)
        docx_path = _convert_doc_to_docx_win32(file_path)
        if docx_path:
            result = index_docx(collection_name, docx_path, filename)
            try:
                import shutil
                shutil.rmtree(os.path.dirname(docx_path), ignore_errors=True)
            except Exception:
                pass
            return result
        # Fallback: extract binary text and index as plain text chunks
        print(f"[DOC] Word COM not available, using binary text fallback for '{filename}'")
        raw_text = _extract_doc_text_fallback(file_path)
        if not raw_text.strip():
            return {"ok": False, "error": "Could not extract text from .doc file. Please save as .docx or .pdf and re-upload."}
        doc_id = str(uuid.uuid4())
        chunks = chunk_text(raw_text, chunk_size=2000, overlap=140)
        stats = {"type": "doc_legacy", "mode": "binary_fallback", "chars": len(raw_text)}
        return _index_text_units(collection_name, doc_id, filename, chunks, doc_type="docx", extra_stats=stats)
    if ext == ".xlsx":
        return index_xlsx(collection_name, file_path, filename)
    if ext == ".csv":
        return index_csv(collection_name, file_path, filename)

    return {"ok": False, "error": f"Unsupported document type: {ext}. Supported: PDF, DOCX, DOC, XLSX, CSV."}


# -------------------------------------------------------------------
# Vector search helper
# -------------------------------------------------------------------
def _vector_search(
    col: str, query: str, top_k: int = 20, doc_ids: Optional[List[str]] = None
) -> List[Tuple[str, dict]]:
    qvec = ollama_embed(query, model=EMBED_MODEL)
    results: List[Tuple[str, dict]] = []
    collection = _get_col(col)

    if not doc_ids:
        res = collection.query(
            query_embeddings=[qvec],
            n_results=top_k,
            include=["documents", "metadatas", "distances"],
        )
        docs = (res.get("documents") or [[]])[0]
        metas = (res.get("metadatas") or [[]])[0]
        results = list(zip(docs, metas))
    else:
        per_doc_k = max(3, top_k // max(1, len(doc_ids)))
        for did in doc_ids:
            r = collection.query(
                query_embeddings=[qvec],
                n_results=per_doc_k,
                where={"doc_id": did},
                include=["documents", "metadatas", "distances"],
            )
            docs = (r.get("documents") or [[]])[0]
            metas = (r.get("metadatas") or [[]])[0]
            results.extend(zip(docs, metas))

    return results[:top_k]


# -------------------------------------------------------------------
# Hybrid Retrieval Pipeline (Vector + BM25 + Multi-Query + Re-rank)
# -------------------------------------------------------------------
def _hybrid_retrieve(
    col: str,
    question: str,
    doc_ids: Optional[List[str]] = None,
    top_k: int = 12,
) -> List[Tuple[str, dict]]:
    retrieve_k = top_k * 3

    queries = [question]
    if ENABLE_MULTI_QUERY:
        extra_queries = _generate_multi_queries(question)
        queries.extend(extra_queries)
        print(f"[RAG:{col}] Searching with {len(queries)} queries (original + {len(extra_queries)} variations)")

    all_vector_results: List[Tuple[str, dict]] = []
    all_bm25_results: List[Tuple[str, dict]] = []

    for q in queries:
        vr = _vector_search(col, q, top_k=retrieve_k, doc_ids=doc_ids)
        all_vector_results.extend(vr)

        if ENABLE_HYBRID_SEARCH:
            br = _bm25_search(col, q, top_k=retrieve_k, doc_ids=doc_ids)
            all_bm25_results.extend(br)

    if ENABLE_HYBRID_SEARCH and all_bm25_results:
        merged = _reciprocal_rank_fusion(all_vector_results, all_bm25_results, top_n=retrieve_k)
        print(f"[RAG:{col}] Hybrid: {len(all_vector_results)} vector + {len(all_bm25_results)} BM25 -> {len(merged)} merged")
    else:
        seen = set()
        merged = []
        for doc, meta in all_vector_results:
            key = f"{meta.get('doc_id', '')}_{meta.get('chunk_index', '')}"
            if key not in seen:
                seen.add(key)
                merged.append((doc, meta))
        merged = merged[:retrieve_k]

    if ENABLE_RERANKING and len(merged) > top_k:
        final = _rerank_with_llm(question, merged, top_n=top_k)
    else:
        final = merged[:top_k]

    return final


# -------------------------------------------------------------------
# SMAC Field-Specific Retrieval (bypasses vector search for known fields)
# -------------------------------------------------------------------
# Common field names found in SMAC tabular PDFs (Col2 values)
SMAC_FIELD_KEYWORDS: List[str] = [
    "gist", "input", "remarks", "narration", "details", "summary",
    "observation", "source", "facts", "allegation", "description",
    "background", "complaint", "incident", "statement", "analysis",
    "intelligence", "information", "findings", "conclusion",
]


def _detect_smac_field(question: str) -> Optional[str]:
    """
    Detect if the question refers to a specific SMAC field name.
    Returns the matched field name (title-cased) or None.
    """
    q_lower = question.lower()
    for field in SMAC_FIELD_KEYWORDS:
        if field in q_lower:
            return field.title()
    return None


def _get_chunks_by_field(
    col: str,
    field_name: str,
    doc_ids: Optional[List[str]] = None,
) -> List[Tuple[str, dict]]:
    """
    Direct ChromaDB metadata lookup for a specific SMAC field — bypasses vector search.
    Tries multiple casings of field_name to handle inconsistent PDF parsing.
    """
    collection = _get_col(col)
    results: List[Tuple[str, dict]] = []
    seen_chunk_ids: set = set()

    # Try common casings: title, lower, UPPER
    casings = {field_name, field_name.lower(), field_name.upper(), field_name.title()}

    for casing in casings:
        try:
            data = collection.get(
                where={"field_name": casing},
                include=["documents", "metadatas"],
            )
            docs = data.get("documents") or []
            metas = data.get("metadatas") or []
            for doc, meta in zip(docs, metas):
                # Filter by doc_ids if specified
                if doc_ids and meta.get("doc_id") not in doc_ids:
                    continue
                # Deduplicate by chunk key
                chunk_key = f"{meta.get('doc_id', '')}_{meta.get('chunk_index', '')}_{doc[:80]}"
                if chunk_key not in seen_chunk_ids:
                    seen_chunk_ids.add(chunk_key)
                    results.append((doc, meta))
        except Exception as e:
            print(f"[FieldRetrieval] Lookup failed for field '{casing}': {e}")

    print(f"[FieldRetrieval] Found {len(results)} chunks for field '{field_name}' in '{col}'")
    return results


# -------------------------------------------------------------------
# IR Field-Specific Retrieval (fuzzy — IR field names are long phrases)
# -------------------------------------------------------------------
# Short keywords that appear inside IR field names (Col2 of Form-16)
IR_FIELD_KEYWORDS: List[str] = [
    # Identity
    "name", "alias", "code name", "code number",
    "nationality", "religion", "caste", "date of birth", "dob", "age",
    "place of birth", "marital", "mother tongue", "language",
    # Organisation & role
    "organization", "position", "links to other",
    # Addresses
    "permanent address", "present address", "previous address",
    # Crime & legal
    "crime number", "cases registered", "surety",
    # Family & associates
    "parent", "family", "relation", "associate", "contact", "helper",
    "advocate", "lawyer", "doctor",
    # Hideout / shelter
    "hideout", "hide out", "safe house", "shelter",
    # Occupation & education
    "occupation", "qualification", "education",
    # Communication
    "mobile", "phone", "landline", "email", "website", "social media",
    "telegram", "facebook", "instagram",
    # Documents
    "passport", "voter", "ration card", "driving license",
    # Finances & property
    "bank", "account", "property", "vehicle", "locker",
    # Physical
    "physical", "height", "weight", "complexion", "tattoo",
    "built", "identification mark", "habit", "spectacle", "manner",
    # Weapons & travel
    "weapon", "arms", "firearm", "explosive", "ammunition", "pistol", "rifle",
    "travel", "abroad", "transport",
    # Training & activities
    "training", "exercise", "practice", "recruitment", "indoctrination",
    # Fundamentalist / intel
    "fundamentalist", "frontal organisation",
    # Miscellaneous
    "motive", "reason", "purpose", "objective", "ideology",
    "radicalization", "radicalisation", "threat",
]


def _detect_ir_field(question: str) -> Optional[str]:
    """
    Detect if the question mentions an IR field keyword.
    Returns the matched keyword (lowercase) or None.
    Sorted by specificity (longer phrases first) to avoid false short matches.
    """
    q_lower = question.lower()
    # Check longer/more-specific keywords first
    for field in sorted(IR_FIELD_KEYWORDS, key=len, reverse=True):
        if field in q_lower:
            return field
    return None


# Abbreviation ↔ full-form mapping for compound IR field names
_IR_ABBREV_MAP: Dict[str, List[str]] = {
    "dob": ["date of birth", "dob"],
    "date of birth": ["date of birth", "dob"],
    "pob": ["place of birth", "pob"],
    "place of birth": ["place of birth", "pob"],
    "age": ["age"],
    "doj": ["date of joining", "doj"],
    "date of joining": ["date of joining", "doj"],
}

# Synonym expansion for IR field keywords used in fuzzy retrieval
# Ensures "associate" also finds "accomplice", "co-accused", etc.
_IR_FIELD_SYNONYMS: Dict[str, List[str]] = {
    "associate":  ["associate", "accomplice", "co-accused", "helper", "companion", "contact", "abettor"],
    "accomplice": ["associate", "accomplice", "co-accused", "helper", "abettor"],
    "helper":     ["helper", "associate", "accomplice", "co-accused"],
    "contact":    ["contact", "associate", "companion"],
    "family":     ["family", "father", "mother", "brother", "sister", "spouse", "wife", "husband", "relation"],
    "weapon":     ["weapon", "arms", "firearm", "explosive", "ammunition", "pistol", "rifle", "equipment"],
    "arms":       ["arms", "weapon", "firearm", "explosive", "ammunition", "equipment"],
    "firearm":    ["firearm", "weapon", "arms", "pistol", "rifle"],
    "training":   ["training", "exercise", "practice", "indoctrination", "recruitment"],
    "motive":     ["motive", "reason", "purpose", "objective", "ideology"],
    "travel":     ["travel", "mode of travel", "arrive", "arrival", "transport", "vehicle"],
    "address":    ["address", "permanent address", "present address", "residence", "village"],
    "phone":      ["mobile", "phone", "landline", "contact number"],
    "bank":       ["bank", "account", "financial"],
    "advocate":   ["advocate", "lawyer", "legal", "counsel"],
    "case":       ["case", "crime number", "fir", "chargesheet", "police station"],
    "visit":      ["visit", "place of visit", "visited"],
    "organization": ["organization", "organisation", "organi", "affiliation", "group", "outfit"],
}


def _get_chunks_by_field_fuzzy(
    col: str,
    keyword: str,
    doc_ids: Optional[List[str]] = None,
) -> List[Tuple[str, dict]]:
    """
    Fuzzy field lookup for IR — scans all chunk metadata and returns chunks
    where field_name contains the keyword (case-insensitive substring match).
    Also handles compound field names split by "/" (e.g., "Age/DOB/Place of Birth")
    and common abbreviations (e.g., "DOB" ↔ "Date of Birth").
    Expands keywords with synonyms (e.g., "associate" also finds "accomplice", "co-accused").
    ChromaDB does not support LIKE queries, so we fetch all and filter in Python.
    """
    collection = _get_col(col)
    all_data = collection.get(include=["documents", "metadatas"])
    docs = all_data.get("documents") or []
    metas = all_data.get("metadatas") or []

    kw_lower = keyword.lower()
    # Build set of all search terms: keyword + abbreviation equivalents + synonyms
    search_terms = {kw_lower}
    if kw_lower in _IR_ABBREV_MAP:
        search_terms.update(t.lower() for t in _IR_ABBREV_MAP[kw_lower])
    if kw_lower in _IR_FIELD_SYNONYMS:
        search_terms.update(t.lower() for t in _IR_FIELD_SYNONYMS[kw_lower])

    field_results: List[Tuple[str, dict]] = []
    text_results: List[Tuple[str, dict]] = []

    for doc, meta in zip(docs, metas):
        if doc_ids and meta.get("doc_id") not in doc_ids:
            continue
        field_name = (meta.get("field_name") or "").lower()

        # Match against field_name metadata
        if field_name:
            if any(term in field_name for term in search_terms):
                field_results.append((doc, meta))
                continue
            # Compound field: split by "/" and check each component
            if "/" in field_name:
                components = [c.strip() for c in field_name.split("/") if c.strip()]
                matched = False
                for comp in components:
                    if any(term in comp or comp in term for term in search_terms):
                        matched = True
                        break
                if matched:
                    field_results.append((doc, meta))
                    continue

        # Also match against document text content (catches page text chunks)
        doc_lower = (doc or "").lower()
        if any(term in doc_lower for term in search_terms):
            text_results.append((doc, meta))

    # Field-name matches first (most precise), then text matches (capped to avoid context flood)
    results = field_results + text_results[:8]
    print(f"[FieldRetrieval-Fuzzy] Found {len(field_results)} field + {min(len(text_results), 8)}/{len(text_results)} text chunks for keyword '{keyword}' (terms: {search_terms}) in '{col}'")
    return results


def _ir_text_search(
    col: str,
    question: str,
    doc_ids: Optional[List[str]] = None,
) -> List[Tuple[str, dict]]:
    """
    Direct text search across all IR chunks — finds chunks containing
    the significant words from the question. Used when no IR field keyword
    is detected, to catch specific document sections that hybrid search misses.
    """
    _q_stop = {
        "what", "who", "where", "when", "how", "which", "is", "are", "was",
        "were", "the", "a", "an", "of", "in", "for", "to", "and", "or",
        "his", "her", "their", "its", "this", "that", "do", "does", "did",
        "has", "have", "had", "be", "been", "being", "will", "can", "could",
        "about", "from", "with", "tell", "me", "give", "show", "find",
        "please", "details", "information", "you", "not", "them", "they",
    }
    words = re.findall(r"[a-zA-Z]{3,}", question.lower())
    search_words = [w for w in words if w not in _q_stop]
    if not search_words:
        return []

    collection = _get_col(col)
    all_data = collection.get(include=["documents", "metadatas"])
    docs = all_data.get("documents") or []
    metas = all_data.get("metadatas") or []

    # Score each chunk by how many search words it contains
    scored: List[Tuple[int, str, dict]] = []
    # For short queries like "hideout places", require at least 1 hit for better recall.
    # For longer queries, require roughly half of the terms.
    if len(search_words) <= 3:
        min_matches = 1
    else:
        min_matches = max(2, (len(search_words) + 1) // 2)

    for doc, meta in zip(docs, metas):
        if doc_ids and meta.get("doc_id") not in doc_ids:
            continue
        doc_lower = (doc or "").lower()
        hits = sum(1 for w in search_words if w in doc_lower)
        if hits >= min_matches:
            scored.append((hits, doc, meta))

    # Sort by match count (best first), keep a larger candidate set for long IR docs.
    scored.sort(key=lambda x: x[0], reverse=True)
    results = [(doc, meta) for _, doc, meta in scored[:25]]

    print(f"[IR-TextSearch] Found {len(results)} chunks (>={min_matches}/{len(search_words)} words matched) for {search_words}")
    return results


# -------------------------------------------------------------------
# Smart routing: detect aggregate/cross-document questions
# -------------------------------------------------------------------
_AGGREGATE_PATTERNS = re.compile(
    r'\b(list all|name all|show all|find all|give all|get all|display all|'
    r'all the|all accused|all subjects|all persons|all documents|all reports|'
    r'every|each document|each person|each accused|each subject|'
    r'how many|count|total number|'
    r'compare|comparison|across all|summary of all|'
    r'tabulate|table of|chart of)\b',
    re.IGNORECASE,
)


def _is_aggregate_question(question: str) -> bool:
    """Detect if a question requires cross-document aggregation (NL→SQL)."""
    return bool(_AGGREGATE_PATTERNS.search(question))


def _answer_via_sql(question: str, collection_name: str) -> Optional[Dict[str, Any]]:
    """
    Answer a question using NL→SQL pipeline against document_fields table.
    Returns {"answer": ..., "sql": ..., "used_chunks": []} or None on failure.
    """
    try:
        schema = get_table_schema_description()

        sql_prompt = (
            "You are a SQL expert for T-SQL (Microsoft SQL Server).\n"
            "Given the database schema below and the user's question, generate a SQL SELECT query.\n\n"
            f"DATABASE SCHEMA:\n{schema}\n\n"
            f"USER QUESTION: {question}\n\n"
            "RULES:\n"
            "- Return ONLY the SQL query, no explanation.\n"
            "- Use T-SQL syntax (TOP instead of LIMIT, NVARCHAR, etc.).\n"
            "- The table is a key-value store: each row has field_key (field name) and field_value.\n"
            "- To find a value, search field_key with LIKE '%keyword%' and return field_value.\n"
            "- Use LIKE with '%keyword%' for text matching.\n"
            f"- Filter by collection = '{collection_name}' to limit to the right document type.\n"
            "- To get multiple fields per document, self-join document_fields on doc_id.\n"
            "- Use LEFT JOIN when one field might be missing (show NULL instead of omitting row).\n"
            "- The name of the accused/criminal/convict/subject is stored with field_key LIKE '%Name%'.\n"
            "  Variations: 'Name', 'Name of the Subject', 'Name of the Accused'.\n"
            "- For helpers use field_key LIKE '%helper%'.\n"
            "- For associates use field_key LIKE '%associate%' OR field_key LIKE '%accomplice%'.\n"
            "- For family use field_key LIKE '%family%' OR field_key LIKE '%father%' etc.\n"
            "- For address use field_key LIKE '%address%'.\n"
            "- For organization use field_key LIKE '%organi%'.\n"
            "- Always use SELECT, never INSERT/UPDATE/DELETE.\n"
            "- Return doc_name for document identification.\n"
            "- When asked for 'all accused' or 'all persons', get the Name field from every document.\n"
        )

        sql_response = ollama_chat(
            [{"role": "user", "content": sql_prompt}],
            temperature=0.0,
            model=PDF_MODEL,
        )

        # Extract SQL (strip markdown fences)
        sql_query = sql_response.strip()
        if sql_query.startswith("```"):
            lines = sql_query.split("\n")
            sql_lines = [ln for ln in lines if not ln.strip().startswith("```")]
            sql_query = "\n".join(sql_lines).strip()
        if sql_query.lower().startswith("sql"):
            sql_query = sql_query[3:].strip()

        print(f"[RAG:NL→SQL] Generated query: {sql_query}")

        rows = execute_sql_query(sql_query)

        # Check for errors
        if rows and len(rows) == 1 and "error" in rows[0]:
            print(f"[RAG:NL→SQL] SQL error: {rows[0]['error']}")
            return None

        if not rows:
            return None

        # Format results with LLM
        import json
        result_text = json.dumps(rows[:50], indent=2, default=str)

        answer_prompt = (
            "You are an authorized internal AI assistant for Karnataka State Police (KSP).\n"
            "ALWAYS respond in English only.\n"
            "Given the user's question and the SQL query results, provide a clear, "
            "factual answer based ONLY on the data returned.\n\n"
            f"USER QUESTION: {question}\n\n"
            f"SQL RESULTS ({len(rows)} rows):\n{result_text}\n\n"
            "RULES:\n"
            "- Answer in English only, regardless of the language of the data.\n"
            "- Answer in clear, readable format.\n"
            "- Present tabular data as a formatted list or table.\n"
            "- If a field is NULL or missing for a person, show it as 'Not mentioned' or '-'.\n"
            "- Include all relevant data from the results.\n"
            "- Do NOT add information not present in the results.\n"
        )

        answer = ollama_chat(
            [{"role": "user", "content": answer_prompt}],
            temperature=0.0,
            model=PDF_MODEL,
        )

        print(f"[RAG:NL→SQL] Answer generated from {len(rows)} SQL rows")
        return {"answer": answer, "sql": sql_query, "used_chunks": []}

    except Exception as e:
        print(f"[RAG:NL→SQL] Failed: {e}")
        return None


# -------------------------------------------------------------------
# Hallucination guard: verify LLM answer is grounded in context
# -------------------------------------------------------------------
_GROUNDING_STOP_WORDS = {
    "the", "a", "an", "is", "are", "was", "were", "be", "been", "being",
    "have", "has", "had", "do", "does", "did", "will", "would", "could",
    "should", "may", "might", "shall", "can", "need", "must",
    "and", "but", "or", "nor", "not", "so", "yet", "both", "either",
    "for", "with", "from", "into", "about", "between", "through",
    "this", "that", "these", "those", "which", "what", "who", "whom",
    "there", "here", "where", "when", "how", "why", "then", "than",
    "very", "also", "just", "only", "more", "most", "some", "any",
    "each", "every", "such", "other", "another", "same",
    "based", "according", "mentioned", "document", "documents",
    "provided", "above", "below", "found", "information", "context",
    "following", "however", "therefore", "regarding", "report",
}


def _verify_grounding(answer: str, context: str, question: str) -> str:
    """
    Post-processing hallucination guard: checks each sentence in the LLM answer
    against the context. Sentences with very low word overlap with the context
    are likely hallucinated from training data and are removed.
    """
    if not answer or not answer.strip():
        return answer

    # Build a set of significant words from the context (lowercase, >= 4 chars)
    context_lower = context.lower()
    context_words = set(re.findall(r"[a-zA-Z]{4,}", context_lower))

    # Also include question words as valid (the answer may echo the question)
    q_words = set(re.findall(r"[a-zA-Z]{4,}", question.lower()))
    valid_words = context_words | q_words | _GROUNDING_STOP_WORDS

    # Split answer into sentences (handle numbered lists, bullet points, newlines)
    sentences = re.split(r'(?<=[.!?])\s+|\n+', answer.strip())

    kept = []
    removed = []
    for sent in sentences:
        sent = sent.strip()
        if not sent:
            continue

        # Extract significant words from this sentence
        sent_words = [w for w in re.findall(r"[a-zA-Z]{4,}", sent.lower())
                       if w not in _GROUNDING_STOP_WORDS]

        if not sent_words:
            # Sentence has only short/stop words — keep it (e.g., "Yes.", "No.")
            kept.append(sent)
            continue

        # Calculate grounding ratio: what fraction of the sentence's words appear in context?
        grounded_count = sum(1 for w in sent_words if w in valid_words)
        ratio = grounded_count / len(sent_words)

        # Determine threshold based on sentence type:
        # - List items (start with number/bullet) → lenient 0.20 (names may not repeat verbatim)
        # - Short sentences ≤ 60 chars → standard 0.35 (hallucinations like "Corkage fees..." must be caught)
        # - Long sentences → standard 0.35
        is_list_item = bool(re.match(r'^[\d]+[.)]\s|^[-\*\•\–]\s', sent))
        threshold = 0.20 if is_list_item else 0.35

        if ratio >= threshold:
            kept.append(sent)
        else:
            removed.append((sent, f"{ratio:.0%}"))

    if removed:
        print(f"[RAG:Grounding] Removed {len(removed)} ungrounded sentence(s):")
        for sent, pct in removed:
            print(f"  - ({pct} grounded) {sent[:120]}")

    if not kept:
        return "This information is not found in the provided documents."

    return "\n".join(kept)


def _strip_non_latin(text: str) -> str:
    """
    Remove non-Latin script characters (Kannada, Devanagari, Tamil, etc.)
    from LLM answers to enforce English-only output.
    Keeps: Latin letters, digits, punctuation, common symbols, whitespace.
    """
    # Match runs of non-Latin script (Unicode blocks for Indian languages, CJK, Arabic, etc.)
    # Keep: Basic Latin, Latin Extended, digits, punctuation, symbols, whitespace
    cleaned = re.sub(r'[^\x00-\x7F\u00C0-\u024F]+', '', text)
    # Clean up leftover artifacts: multiple spaces, empty parentheses, dangling punctuation
    cleaned = re.sub(r'\(\s*\)', '', cleaned)       # empty ()
    cleaned = re.sub(r'\s{2,}', ' ', cleaned)       # collapse multiple spaces
    cleaned = re.sub(r'^\s*[,;:]\s*', '', cleaned, flags=re.MULTILINE)  # leading punctuation on lines
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)    # collapse blank lines
    if cleaned.strip() != text.strip():
        print(f"[RAG] Stripped non-Latin characters from answer ({len(text)} → {len(cleaned)} chars)")
    return cleaned.strip()


# -------------------------------------------------------------------
# Ask across documents
# -------------------------------------------------------------------
def ask_docs(question: str, doc_ids: Optional[List[str]] = None, top_k: int = 12, collection_name: str = "SMAC", raw_question: Optional[str] = None) -> Dict[str, Any]:
    # Use raw_question for keyword extraction/field detection (avoids conversation history pollution)
    # Fall back to full question if raw_question is not provided
    q_for_keywords = raw_question if raw_question else question
    print(f"[RAG] raw_question={repr(raw_question[:80]) if raw_question else 'None'}, q_for_keywords={repr(q_for_keywords[:80])}")

    # ── Smart routing: aggregate/cross-document questions → NL→SQL ─────
    if _is_aggregate_question(q_for_keywords):
        print(f"[RAG] Aggregate question detected — routing through NL→SQL pipeline")
        sql_result = _answer_via_sql(question, collection_name)
        if sql_result:
            return sql_result
        print(f"[RAG] NL→SQL returned no results — falling back to hybrid approach")

    # For SMAC: detect if the question targets a specific field and use direct metadata lookup
    if collection_name == "SMAC":
        detected_field = _detect_smac_field(q_for_keywords)
        if detected_field:
            print(f"[RAG:SMAC] Field-specific query detected: '{detected_field}' — bypassing vector search")
            field_chunks = _get_chunks_by_field(collection_name, detected_field, doc_ids)
            if field_chunks:
                results = field_chunks[:top_k]
            else:
                # Field not found in metadata — fall back to hybrid retrieval
                print(f"[RAG:SMAC] No chunks found for field '{detected_field}', falling back to hybrid retrieval")
                results = _hybrid_retrieve(collection_name, q_for_keywords, doc_ids=doc_ids, top_k=top_k)
        else:
            results = _hybrid_retrieve(collection_name, q_for_keywords, doc_ids=doc_ids, top_k=top_k)

    elif collection_name == "IR":
        detected_keyword = _detect_ir_field(q_for_keywords)
        if detected_keyword:
            print(f"[RAG:IR] Field keyword detected: '{detected_keyword}' — using fuzzy field retrieval")
            field_chunks = _get_chunks_by_field_fuzzy(collection_name, detected_keyword, doc_ids)
        else:
            # No known field keyword — do a direct text search for question words
            field_chunks = _ir_text_search(collection_name, q_for_keywords, doc_ids)

        # Merge field-specific results with hybrid retrieval
        # Field/text matches are the primary answer — include ALL of them
        # Hybrid results are supplementary context — cap those to reduce noise
        if field_chunks:
            hybrid_cap = min(5, top_k)
            hybrid_chunks = _hybrid_retrieve(collection_name, q_for_keywords, doc_ids=doc_ids, top_k=hybrid_cap)
        else:
            hybrid_chunks = _hybrid_retrieve(collection_name, q_for_keywords, doc_ids=doc_ids, top_k=top_k)

        # Merge: field/text-matched chunks first (most targeted), then hybrid
        seen_keys = set()
        results = []
        for d, m in (field_chunks or []):
            key = f"{m.get('doc_id', '')}_{m.get('chunk_index', '')}"
            if key not in seen_keys:
                seen_keys.add(key)
                results.append((d, m))
        for d, m in hybrid_chunks:
            key = f"{m.get('doc_id', '')}_{m.get('chunk_index', '')}"
            if key not in seen_keys:
                seen_keys.add(key)
                results.append((d, m))
        print(f"[RAG:IR] Merged: {len(field_chunks or [])} field/text + {len(hybrid_chunks)} hybrid → {len(results)} results (no field cap)")

    else:
        results = _hybrid_retrieve(collection_name, q_for_keywords, doc_ids=doc_ids, top_k=top_k)

    # ── Structured table keyword search ──────────────────────────────────
    # Search document_fields for matching field_keys and add as extra context
    _stop_words = {
        "what", "who", "where", "when", "how", "which", "is", "are", "was",
        "were", "the", "a", "an", "of", "in", "for", "to", "and", "or",
        "his", "her", "their", "its", "this", "that", "these", "those",
        "do", "does", "did", "has", "have", "had", "be", "been", "being",
        "will", "can", "could", "should", "would", "may", "might",
        "about", "from", "with", "tell", "me", "give", "show", "find",
        "list", "all", "any", "some", "please", "details", "information",
        # Conversation template and history noise words
        "context", "conversation", "continuing", "ongoing", "resolve",
        "references", "previous", "answer", "question", "section",
        "assistant", "user", "current", "provided", "document", "documents",
        "same", "like", "use", "not", "you", "him", "them", "they",
        "also", "just", "get", "got", "said", "yes", "know", "found",
        "getting", "helped", "introduced",
    }
    # Synonym expansion: user terms → field_key search terms
    _KEYWORD_SYNONYMS = {
        "accused": ["name", "accused"],
        "criminal": ["name", "criminal"],
        "convict": ["name", "convict"],
        "subject": ["name", "subject"],
        "person": ["name", "person"],
        "suspect": ["name", "suspect"],
        "perpetrator": ["name", "perpetrator"],
        "offender": ["name", "offender"],
        "address": ["address", "permanent address", "present address"],
        "phone": ["mobile", "landline", "phone"],
        "family": ["father", "mother", "brother", "sister", "spouse", "family"],
        "weapon": ["weapon", "arms"],
        "travel": ["travel", "mode of travel", "arrive", "arrival"],
        "arrive": ["travel", "mode of travel", "arrive", "arrival"],
        "arrival": ["travel", "mode of travel", "arrive", "arrival"],
        "organization": ["organi", "affiliation"],
        "organisation": ["organi", "affiliation"],
        "associate": ["associate", "accomplice", "helper", "co-accused", "companion", "contact"],
        "accomplice": ["associate", "accomplice", "helper", "co-accused"],
        "helper": ["associate", "helper", "co-accused"],
        "companion": ["associate", "companion", "contact"],
        "hideout": ["hideout", "hide out", "safe house", "shelter", "place of hideout"],
        "hideouts": ["hideout", "hide out", "safe house", "shelter", "place of hideout"],
        "safehouse": ["safe house", "hideout", "hide out", "shelter"],
        "shelter": ["shelter", "hideout", "safe house"],
        "lawyer": ["advocate", "lawyer", "legal", "counsel"],
        "advocate": ["advocate", "lawyer", "legal", "counsel"],
        "counsel": ["advocate", "lawyer", "counsel"],
    }
    structured_context_blocks: List[str] = []
    try:
        words = re.findall(r"[a-zA-Z]{3,}", q_for_keywords.lower())
        keywords = [w for w in words if w not in _stop_words]
        # Expand with synonyms
        expanded = set(keywords)
        for kw in keywords:
            if kw in _KEYWORD_SYNONYMS:
                expanded.update(_KEYWORD_SYNONYMS[kw])
        keywords = list(expanded)
        seen_keys: set = set()
        for kw in keywords:
            matches = search_fields(kw, collection=collection_name)
            for match in matches:
                # Keep repeated rows with the same field_key but different values (common in IR).
                dedup_key = (
                    match["doc_id"],
                    match["field_key"],
                    (match.get("serial_no") or ""),
                    (match.get("field_value") or ""),
                )
                if dedup_key not in seen_keys:
                    seen_keys.add(dedup_key)
                    structured_context_blocks.append(
                        f"[STRUCTURED DATA | {match['doc_name']} | Sl {match.get('serial_no', '?')}]\n"
                        f"{match['field_key']}: {match['field_value']}"
                    )
        if structured_context_blocks:
            print(f"[RAG] Structured table search found {len(structured_context_blocks)} matching fields for keywords: {keywords}")
    except Exception as e:
        print(f"[RAG] Structured table search failed: {e}")

    used: List[dict] = []
    context_blocks: List[str] = structured_context_blocks.copy()

    for d, m in results:
        used.append({
            "doc_id": m.get("doc_id"),
            "doc_name": m.get("doc_name"),
            "doc_type": m.get("doc_type"),
            "page": m.get("page"),
            "sheet": m.get("sheet"),
            "chunk_index": m.get("chunk_index"),
        })

        loc = []
        if m.get("page") is not None:
            loc.append(f"page {m.get('page')}")
        if m.get("sheet"):
            loc.append(f"sheet {m.get('sheet')}")
        loc_str = (", ".join(loc)) if loc else "location ?"

        context_blocks.append(
            f"[{m.get('doc_name')} | {m.get('doc_type')} | {loc_str} | chunk {m.get('chunk_index')}]\n{d}"
        )

    context = "\n\n".join(context_blocks).strip()
    if not context:
        return {"answer": "Not found in the document(s).", "used_chunks": []}

    # Cap context — gemma3:12b supports 128K tokens (~400K chars)
    # List queries need more context to find all entries
    _list_detect = re.compile(
        r'\b(list all|list the|who are|who were|names of|name all|give all|show all|'
        r'all the|how many|enumerate|what are all|tell me all)\b',
        re.IGNORECASE,
    )
    MAX_CONTEXT_CHARS = 60000 if _list_detect.search(q_for_keywords) else 30000
    if len(context) > MAX_CONTEXT_CHARS:
        print(f"[RAG] Context trimmed: {len(context)} → {MAX_CONTEXT_CHARS} chars")
        context = context[:MAX_CONTEXT_CHARS]

    is_summarize = "summar" in question.lower()

    # Detect list-type questions — needs exhaustive extraction, not just top-match
    _list_patterns = re.compile(
        r'\b(list all|list the|who are|who were|names of|name all|give all|show all|'
        r'all the|how many|enumerate|what are all|tell me all)\b',
        re.IGNORECASE,
    )
    is_list_query = bool(_list_patterns.search(llm_question if q_for_keywords != question else question))

    # Use raw question for the LLM prompt (avoids sending conversation history as noise)
    llm_question = q_for_keywords if q_for_keywords != question else question

    if is_list_query:
        print(f"[RAG] List-type query detected — using exhaustive list extraction prompt")
        prompt = (
            "STRICT INSTRUCTION: Answer ONLY from the CONTEXT below. Do NOT use your training knowledge.\n"
            "ALWAYS respond in English only, regardless of the language of the source documents.\n\n"
            "TASK: Extract a COMPLETE numbered list of every entry relevant to the question.\n"
            "Instructions:\n"
            "- Scan ALL sections and ALL context blocks below — do not stop early.\n"
            "- Look for patterns like (i), (ii), (iii), (1), (2), (3), numbered entries, and repeated label groups.\n"
            "- Each unique person/item must appear as a separate numbered entry.\n"
            "- Include the name and any associated details (role, address, relationship) found in the same context block.\n"
            "- Do NOT merge different people into one entry.\n"
            "- Do NOT skip entries because they seem similar.\n"
            "- Do NOT use your training data — only use what is explicitly written in the CONTEXT.\n"
            "- If an entry has no name but has a role/description, include that.\n\n"
            f"CONTEXT:\n{context}\n\n"
            f"QUESTION:\n{llm_question}\n\n"
            "Output a clean numbered list. Include ALL entries found. Do not truncate."
        )
    else:
        prompt = (
            "STRICT INSTRUCTION: Answer ONLY from the CONTEXT below. Do NOT use your training knowledge.\n"
            "ALWAYS respond in English only, regardless of the language of the source documents.\n"
            "If a fact is not explicitly stated in the CONTEXT, respond: 'This information is not found in the provided documents.'\n\n"
            "Rules:\n"
            "- Always answer in English.\n"
            "- Labels may be messy — search for nearby key/value pairs to find requested fields.\n"
            "- Return found fields; mark any missing fields as 'Not found in documents'.\n"
            "- Be EXHAUSTIVE: list ALL names, items, and entries found in the CONTEXT. Never truncate or abbreviate lists.\n"
            "- For summaries: summarize ONLY what is present in the CONTEXT — do not add background knowledge.\n"
            "- Do NOT guess, invent, infer, or assume any fact not explicitly present in the CONTEXT.\n"
            "- Do NOT use your training data to fill gaps — if it is not in the CONTEXT, it does not exist.\n\n"
            f"CONTEXT:\n{context}\n\n"
            f"QUESTION:\n{llm_question}\n\n"
            "FINAL REMINDER: Every fact in your answer must come directly from the CONTEXT above. "
            "If not found there, say so explicitly."
        )

    print(f"[RAG] Sending to LLM: context={len(context)} chars, question={len(llm_question)} chars, prompt={len(prompt)} chars")

    try:
        answer = ollama_chat(
            [
                {"role": "system", "content": (
                    "You are an authorized internal AI assistant for Karnataka State Police (KSP), "
                    "deployed on a secure offline government system for police officers. "
                    "ALWAYS respond in English only, regardless of the language of the source documents. "
                    "CRITICAL RULE: You MUST answer using ONLY the CONTEXT provided in the user message. "
                    "Your training knowledge is completely irrelevant — do NOT use it under any circumstances. "
                    "If the answer is not in the CONTEXT, say exactly: 'This information is not found in the provided documents.' "
                    "Do NOT infer, extrapolate, or fill gaps from your training. "
                    "Do NOT refuse to share information that IS present in the CONTEXT — share it fully and factually."
                )},
                {"role": "user", "content": prompt},
            ],
            temperature=0.0,
            model=PDF_MODEL,
        )
        print(f"[RAG] LLM response received: {len(answer)} chars")
    except Exception as e:
        print(f"[RAG] LLM call FAILED: {e}")
        return {"answer": f"Error generating answer: {e}", "used_chunks": used}

    if not is_summarize and len(answer) > 8000:
        answer = answer[:8000] + "..."

    # ── Hallucination guard: verify answer is grounded in context ──────
    answer = _verify_grounding(answer, context, llm_question)

    # ── Strip non-Latin characters (Kannada, Hindi, etc.) — enforce English-only output ──
    answer = _strip_non_latin(answer)

    return {"answer": answer, "used_chunks": used}


# -------------------------------------------------------------------
# Agentic multi-doc Q&A (find common entities across documents)
# -------------------------------------------------------------------
def ask_docs_agent(question: str, doc_ids: List[str], top_k: int = 6, collection_name: str = "SMAC") -> Dict[str, Any]:
    if not doc_ids or len(doc_ids) < 2:
        return {"answer": "Please select at least two documents.", "used_chunks": []}

    qvec = ollama_embed(question, model=EMBED_MODEL)
    collection = _get_col(collection_name)

    per_doc_context: List[str] = []
    used: List[dict] = []

    for did in doc_ids:
        r = collection.query(
            query_embeddings=[qvec],
            n_results=max(3, top_k),
            where={"doc_id": did},
            include=["documents", "metadatas", "distances"],
        )
        docs = (r.get("documents") or [[]])[0]
        metas = (r.get("metadatas") or [[]])[0]

        if not docs:
            continue

        blocks = []
        for d, m in zip(docs, metas):
            used.append({
                "doc_id": m.get("doc_id"),
                "doc_name": m.get("doc_name"),
                "doc_type": m.get("doc_type"),
                "page": m.get("page"),
                "sheet": m.get("sheet"),
                "chunk_index": m.get("chunk_index"),
            })

            loc = []
            if m.get("page") is not None:
                loc.append(f"page {m.get('page')}")
            if m.get("sheet"):
                loc.append(f"sheet {m.get('sheet')}")
            loc_str = (", ".join(loc)) if loc else "location ?"

            blocks.append(
                f"[{m.get('doc_name')} | {m.get('doc_type')} | {loc_str} | chunk {m.get('chunk_index')}]\n{d}"
            )

        per_doc_context.append("\n\n".join(blocks))

    context = "\n\n---\n\n".join(per_doc_context).strip()
    if not context:
        return {"answer": "Not found in the selected documents.", "used_chunks": []}

    prompt = (
        "You are an evidence-based assistant for Law Enforcement documents.\n"
        "Task: find information that is COMMON across the selected documents.\n"
        "Be generic: compare ANY fields/labels present in the context (names, associates, "
        "accounts, phone numbers, addresses, IDs, dates, etc.).\n"
        "Do NOT assume specific fields exist. Only use what appears in the evidence.\n"
        "Use ONLY the provided context. Do not guess.\n\n"
        "Output format:\n"
        "- Common items list (group by field if possible)\n"
        "- Evidence per document (brief)\n\n"
        f"CONTEXT:\n{context}\n\n"
        f"QUESTION:\n{question}\n"
    )

    answer = ollama_chat(
        [
            {"role": "system", "content": (
                "You are an authorized internal AI assistant for Karnataka State Police (KSP). "
                "You are deployed on a secure, offline government system used exclusively by police officers "
                "for case investigation and analysis. "
                "You MUST answer questions about the provided documents fully and factually. "
                "The documents are official case files and evidence — you are authorized and required to discuss them. "
                "Do NOT refuse, censor, or redact any information from the provided context. "
                "Answer strictly from the evidence/context provided. Do not guess or invent facts."
            )},
            {"role": "user", "content": prompt},
        ],
        temperature=0.0,
        model=PDF_MODEL,
    )

    return {"answer": answer, "used_chunks": used}


def ask_pdf(doc_id: Optional[str], question: str, top_k: int = 12, collection_name: str = "SMAC", raw_question: Optional[str] = None) -> Dict[str, Any]:
    if doc_id:
        return ask_docs(question=question, doc_ids=[doc_id], top_k=top_k, collection_name=collection_name, raw_question=raw_question)
    return ask_docs(question=question, doc_ids=None, top_k=max(top_k, 10), collection_name=collection_name, raw_question=raw_question)


def ask_doc(doc_id: Optional[str], question: str, top_k: int = 10, collection_name: str = "SMAC", raw_question: Optional[str] = None) -> Dict[str, Any]:
    if doc_id:
        return ask_docs(question=question, doc_ids=[doc_id], top_k=top_k, collection_name=collection_name, raw_question=raw_question)
    return ask_docs(question=question, doc_ids=None, top_k=max(top_k, 10), collection_name=collection_name, raw_question=raw_question)


# -------------------------------------------------------------------
# Get all doc chunks grouped by doc_id
# -------------------------------------------------------------------
def get_indexed_doc_list(collection_name: str = "SMAC") -> List[Dict[str, Any]]:
    try:
        collection = _get_col(collection_name)
        count = collection.count()
        if count == 0:
            return []

        all_data = collection.get(include=["metadatas"])
        metas_list = all_data.get("metadatas") or []

        grouped: Dict[str, Dict[str, Any]] = {}
        for meta in metas_list:
            did = meta.get("doc_id", "")
            if not did:
                continue
            if did not in grouped:
                grouped[did] = {
                    "doc_id": did,
                    "doc_name": meta.get("doc_name", "Unknown"),
                    "chunks": 0,
                }
            grouped[did]["chunks"] += 1

        return list(grouped.values())
    except Exception as e:
        print(f"[RAG] get_indexed_doc_list failed for '{collection_name}': {e}")
        return []


def get_all_doc_chunks(collection_name: str = "SMAC", field_filter: Optional[List[str]] = None) -> List[Dict[str, Any]]:
    """
    Return all doc chunks grouped by doc_id.
    field_filter: if provided (e.g. ["Gist", "Input"]), only return chunks whose
                  metadata field_name matches one of the values (case-insensitive).
                  Chunks without a field_name are excluded when field_filter is set.
    """
    try:
        collection = _get_col(collection_name)
        count = collection.count()
        if count == 0:
            return []

        all_data = collection.get(include=["documents", "metadatas"])
        docs_list = all_data.get("documents") or []
        metas_list = all_data.get("metadatas") or []

        normalized_filter = {f.lower() for f in field_filter} if field_filter else None

        grouped: Dict[str, Dict[str, Any]] = {}
        for doc_text, meta in zip(docs_list, metas_list):
            did = meta.get("doc_id", "")
            if not did:
                continue
            if normalized_filter:
                field_name = (meta.get("field_name") or "").lower()
                if field_name not in normalized_filter:
                    continue
            if did not in grouped:
                grouped[did] = {
                    "doc_id": did,
                    "doc_name": meta.get("doc_name", "Unknown"),
                    "chunks": [],
                }
            grouped[did]["chunks"].append(doc_text)

        return list(grouped.values())
    except Exception as e:
        print(f"[RAG] get_all_doc_chunks failed for '{collection_name}': {e}")
        return []


# -------------------------------------------------------------------
# Clear all documents in a named collection
# -------------------------------------------------------------------
def clear_all_documents(collection_name: str = "SMAC") -> Dict[str, Any]:
    global _col_cache
    chroma_name = _to_chroma_name(collection_name)
    try:
        try:
            _client.delete_collection(name=chroma_name)
        except Exception:
            pass

        _col_cache[collection_name] = _client.get_or_create_collection(name=chroma_name)
        _clear_bm25(collection_name)
        return {"ok": True, "message": f"All documents cleared from '{collection_name}'."}

    except Exception as e:
        return {"ok": False, "error": str(e)}
